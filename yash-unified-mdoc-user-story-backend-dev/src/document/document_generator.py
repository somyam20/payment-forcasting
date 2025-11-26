import os
import base64
from datetime import datetime
from typing import List, Tuple, Dict, Any, Optional
from io import BytesIO
import re
import tempfile
import requests
import time
import subprocess
# Document generation libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black, blue, red, green
import litellm
from litellm import completion


# For diagram generation using mermaid.ink API
from PIL import Image as PILImage
import base64
import urllib.parse
from io import BytesIO
from reportlab.platypus import Paragraph, Spacer, Image
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from PIL import Image as PILImage
from reportlab.platypus import Preformatted
import urllib.parse
import os
import logging
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import streamlit as st

import logging
# For OpenAI integration
import json
import os

# Import our centralized OpenAI configuration
from ..utils.openai_config import (
    get_openai_client,
    get_chat_model_name,
    OPENAI_AVAILABLE,
    USE_AZURE,
    AZURE_OPENAI_ENDPOINT,
    AZURE_OPENAI_API_KEY,
    AZURE_OPENAI_API_VERSION,
    OPENAI_API_KEY
)
import logging

from ..utils.logger_config import setup_logger

setup_logger()

import logging
import requests
from io import BytesIO
from PIL import Image as PILImage
from openai import AzureOpenAI
import json

from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from PIL import Image as PILImage
from reportlab.platypus import Preformatted
import base64
import urllib.parse
import os
import logging
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

# For OpenAI integration
import json
import os
import base64
from datetime import datetime
from typing import List, Tuple, Dict, Any, Optional
from io import BytesIO
import re
import tempfile
import requests
import time

# Document generation libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black, blue, red, green

# For diagram generation using mermaid.ink API
from PIL import Image as PILImage
import base64
import urllib.parse
from io import BytesIO
from reportlab.platypus import Paragraph, Spacer, Image
# Import our centralized OpenAI configuration
from ..utils.openai_config import get_openai_client, get_chat_model_name, OPENAI_AVAILABLE, USE_AZURE
import logging

from ..utils.logger_config import setup_logger
from src.utils.obs import LLMUsageTracker

setup_logger()

import base64
import time
import logging
import requests
from io import BytesIO
from PIL import Image as PILImage
import json
import subprocess
import requests
from io import BytesIO
from PIL import Image as PILImage
import graphviz
import re
from ..utils.api_usage_logger import log_openai_usage, log_whisper_usage

token_tracker = LLMUsageTracker()

class MermaidDiagramGenerator:
    def __init__(self):
        """
        Initialize the Enhanced Mermaid Diagram Generator with Azure OpenAI credentials.
        """
        self.azure_client = get_openai_client()
        self.fallback_apis = [
            "https://mermaid.ink/img/",
            "https://kroki.io/mermaid/svg/",
            "https://quickchart.io/mermaid?chart="
        ]
    
    def _convert_mermaid_to_dot(self, mermaid_code: str) -> str:
        """
        Convert Mermaid flowchart to DOT format using AI.
        
        Args:
            mermaid_code: The mermaid code to convert
            
        Returns:
            DOT format code or None if conversion fails
        """
        try:
            logging.info("DEBUG: Converting Mermaid to DOT format using AI")
            
            system_prompt = """You are an expert in both Mermaid and Graphviz DOT syntax. Your task is to convert Mermaid flowcharts to equivalent DOT format.

Rules:
1. Only return the DOT code, nothing else
2. Maintain the original structure and flow as much as possible
3. Use appropriate DOT syntax for nodes, edges, and styling
4. Handle different Mermaid node shapes by using DOT shape attributes
5. Preserve labels and text content
6. Use proper DOT syntax for directed graphs

Common conversions:
- flowchart TD/TB -> digraph { rankdir=TB; }
- flowchart LR -> digraph { rankdir=LR; }
- A --> B -> A -> B;
- A[label] -> A [label="label"];
- A((circle)) -> A [shape=circle, label="circle"];
- A{diamond} -> A [shape=diamond, label="diamond"];
- A[/parallelogram/] -> A [shape=parallelogram, label="parallelogram"];

Example:
Mermaid: flowchart TD; A[Start] --> B{Decision}; B --> C[End];
DOT: digraph { rankdir=TB; A [label="Start"]; B [shape=diamond, label="Decision"]; C [label="End"]; A -> B; B -> C; }"""

            user_prompt = f"""Convert this Mermaid flowchart to DOT format:

```
{mermaid_code}
```

Please provide only the DOT code."""

            response = self.azure_client.chat.completions.create(
                model=get_chat_model_name(),
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=1500
            )
            
            # Log usage - this is the correct way to access usage
            usage = response.usage
            if usage:
                log_openai_usage(
                    module_name=__name__,
                    model=get_chat_model_name(),
                    prompt_tokens=usage.prompt_tokens,
                    completion_tokens=usage.completion_tokens,
                    function_name="_convert_mermaid_to_dot"
                )
            
            dot_code = response.choices[0].message.content.strip()
            
            # Clean up the response - remove markdown code blocks if present
            if dot_code.startswith("```"):
                lines = dot_code.split('\n')
                if lines[0].startswith("```") and lines[-1].strip() == "```":
                    dot_code = '\n'.join(lines[1:-1])
                elif lines[0].startswith("```"):
                    dot_code = '\n'.join(lines[1:])
            
            logging.info(f"DEBUG: AI converted to DOT: {dot_code[:100]}...")
            return dot_code.strip()
            
        except Exception as e:
            logging.error(f"Failed to convert Mermaid to DOT: {e}")
            return None
    
    def _convert_mermaid_to_plantuml(self, mermaid_code: str) -> str:
        """
        Convert Mermaid to PlantUML using AI.
        
        Args:
            mermaid_code: The mermaid code to convert
            
        Returns:
            PlantUML code or None if conversion fails
        """
        try:
            system_prompt = """Convert Mermaid flowcharts to PlantUML syntax. Only return the PlantUML code.

Example:
Mermaid: flowchart TD; A[Start] --> B{Decision}; B --> C[End];
PlantUML: @startuml
start
:Start;
if (Decision) then
:End;
endif
stop
@enduml"""

            response = self.azure_client.chat.completions.create(
                model=get_chat_model_name(),
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Convert to PlantUML:\n{mermaid_code}"}
                ],
                temperature=0.1,
                max_tokens=1000
            )
            
            # Log usage
            usage = response.usage
            if usage:
                log_openai_usage(
                    module_name=__name__,
                    model=get_chat_model_name(),
                    prompt_tokens=usage.prompt_tokens,
                    completion_tokens=usage.completion_tokens,
                    function_name="_convert_mermaid_to_plantuml"
                )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            logging.error(f"PlantUML conversion failed: {e}")
            return None
    
    def _render_with_graphviz(self, mermaid_code: str, dot_code: str = None, attempt_num: int = 1) -> PILImage.Image:
        """
        Render diagram by converting Mermaid to DOT and using Graphviz.
        
        Args:
            mermaid_code: The original mermaid code
            dot_code: Pre-converted DOT code (None to convert from mermaid)
            attempt_num: Current attempt number for logging
            
        Returns:
            PIL Image object, error dict, or None if failed
        """
        try:
            logging.info(f"DEBUG: Graphviz attempt {attempt_num}")
            
            # Convert Mermaid to DOT if not provided
            if dot_code is None:
                dot_code = self._convert_mermaid_to_dot(mermaid_code)
                if not dot_code:
                    return None
            
            # Try to render with Python graphviz library first
            try:
                graph = graphviz.Source(dot_code, format='png')
                png_bytes = graph.pipe()
                img = PILImage.open(BytesIO(png_bytes))
                logging.info(f"DEBUG: Graphviz Python lib success: {img.size}")
                return img
                
            except Exception as py_error:
                logging.warning(f"Graphviz Python library failed: {py_error}")
                
                # Try CLI graphviz as fallback
                try:
                    return self._render_with_graphviz_cli(dot_code)
                except Exception as cli_error:
                    logging.warning(f"Graphviz CLI also failed: {cli_error}")
                    # Return error info for AI fixing
                    return {
                        "error": f"Graphviz errors - Python: {str(py_error)}, CLI: {str(cli_error)}",
                        "dot_code": dot_code,
                        "original_mermaid": mermaid_code
                    }
                
        except Exception as e:
            logging.error(f"Graphviz error: {e}")
            return None
    
    def _render_with_graphviz_cli(self, dot_code: str) -> PILImage.Image:
        """
        Render DOT code using command line Graphviz as fallback.
        
        Args:
            dot_code: The DOT code to render
            
        Returns:
            PIL Image object or None if failed
        """
        try:
            logging.info("DEBUG: Attempting to render with Graphviz CLI")
            
            # Create temporary files
            with tempfile.NamedTemporaryFile(mode='w', suffix='.dot', delete=False) as input_file:
                input_file.write(dot_code)
                input_path = input_file.name
            
            output_path = input_path.replace('.dot', '.png')
            
            try:
                # Run dot command
                result = subprocess.run([
                    'dot', '-Tpng', input_path, '-o', output_path
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0 and os.path.exists(output_path):
                    img = PILImage.open(output_path)
                    logging.info(f"DEBUG: Graphviz CLI success: {img.size}")
                    return img.copy()
                else:
                    logging.warning(f"Graphviz CLI failed: {result.stderr}")
                    return None
                    
            finally:
                # Cleanup temporary files
                for path in [input_path, output_path]:
                    if os.path.exists(path):
                        os.unlink(path)
                        
        except Exception as e:
            logging.error(f"Graphviz CLI error: {e}")
            return None
    
    def _render_with_plantuml_fallback(self, mermaid_code: str) -> PILImage.Image:
        """
        Another fallback: Convert to PlantUML and render.
        
        Args:
            mermaid_code: The mermaid code to render
            
        Returns:
            PIL Image object or None if failed
        """
        try:
            logging.info("DEBUG: Attempting PlantUML fallback")
            
            # Convert Mermaid to PlantUML using AI
            plantuml_code = self._convert_mermaid_to_plantuml(mermaid_code)
            if not plantuml_code:
                return None
            
            # Use PlantUML server to render
            api_url = "http://www.plantuml.com/plantuml/png/"
            
            # Encode for PlantUML
            import zlib
            compressed = zlib.compress(plantuml_code.encode('utf-8'))
            encoded = base64.b64encode(compressed).decode('ascii')
            
            response = requests.get(f"{api_url}{encoded}", timeout=15)
            
            if response.status_code == 200:
                img = PILImage.open(BytesIO(response.content))
                logging.info(f"DEBUG: PlantUML success: {img.size}")
                return img
            
        except Exception as e:
            logging.error(f"PlantUML fallback error: {e}")
        
        return None
    
    def _render_with_kroki(self, mermaid_code: str) -> PILImage.Image:
        """
        Render mermaid diagram using Kroki.io API via POST request.
        
        Args:
            mermaid_code: The mermaid code to render
            
        Returns:
            PIL Image object or None if failed
        """
        try:
            logging.info("DEBUG: Attempting to render with Kroki.io via POST")

            api_url = 'https://kroki.io/mermaid/png'
            headers = {'Content-Type': 'text/plain'}

            response = requests.post(api_url, data=mermaid_code.encode('utf-8'), headers=headers, timeout=15)

            if response.status_code == 200:
                img = PILImage.open(BytesIO(response.content))
                logging.info(f"DEBUG: Kroki.io success: {img.size}")
                return img
            else:
                logging.warning(f"Kroki.io failed: {response.status_code} - {response.text}")
                return None

        except Exception as e:
            logging.error(f"Kroki.io error: {e}")
            return None

    def _render_with_mermaid_cli(self, mermaid_code: str) -> PILImage.Image:
        """
        Render mermaid diagram using local Mermaid CLI (if installed).
        
        Args:
            mermaid_code: The mermaid code to render
            
        Returns:
            PIL Image object or None if failed
        """
        try:
            logging.info("DEBUG: Attempting to render with Mermaid CLI")
            
            # Create temporary files
            with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as input_file:
                input_file.write(mermaid_code)
                input_path = input_file.name
            
            output_path = input_path.replace('.mmd', '.png')
            
            try:
                # Run mermaid CLI
                result = subprocess.run([
                    'npx','mmdc', 
                    '-i', input_path, 
                    '-o', output_path,
                    '-t', 'neutral',
                    '-b', 'white'
                ], capture_output=True, text=True, timeout=30)

                if result.returncode == 0 and os.path.exists(output_path):
                    img = PILImage.open(output_path)
                    logging.info(f"DEBUG: Mermaid CLI success: {img.size}")
                    return img.copy()  # Copy to memory before file cleanup
                else:
                    logging.warning(f"Mermaid CLI failed: {result.stderr}")
                    return None
                    
            finally:
                # Cleanup temporary files
                for path in [input_path, output_path]:
                    if os.path.exists(path):
                        os.unlink(path)
                        
        except Exception as e:
            logging.error(f"Mermaid CLI error: {e}")
            return None
    
    def _get_error_from_response(self, response) -> str:
        """
        Extract error information from the mermaid.ink API response.
        
        Args:
            response: The HTTP response from mermaid.ink
            
        Returns:
            Error message string
        """
        try:
            if response.status_code == 400:
                return f"Bad Request (400): Invalid mermaid syntax. Response: {response.text[:200]}"
            elif response.status_code == 500:
                return f"Server Error (500): Mermaid processing failed. Response: {response.text[:200]}"
            else:
                return f"HTTP {response.status_code}: {response.text[:200]}"
        except:
            return f"HTTP {response.status_code}: Unable to parse error response"
    
    def _render_with_mermaid_ink(self, mermaid_code: str, max_attempts: int = 5) -> PILImage.Image:
        """
        Generate diagram from mermaid code using mermaid.ink with retries.
        
        Args:
            mermaid_code: The mermaid code to render
            max_attempts: Maximum number of attempts
            
        Returns:
            PIL Image object or error dict if failed
        """
        # Clean up the mermaid code
        mermaid_code = mermaid_code.strip()
        
        # Encode the mermaid code for the API
        graphbytes = mermaid_code.encode("utf8")
        base64_bytes = base64.urlsafe_b64encode(graphbytes)
        base64_string = base64_bytes.decode("ascii")
        
        # Get the diagram image from mermaid.ink API
        api_url = f'https://mermaid.ink/img/{base64_string}'

        for attempt in range(max_attempts):
            logging.info(f"DEBUG: Mermaid.ink attempt {attempt + 1} - Calling API: {api_url[:100]}...")
            try:
                response = requests.get(api_url, timeout=15)
                logging.info(f"DEBUG: mermaid.ink API response status: {response.status_code}")

                if response.status_code == 200:
                    img = PILImage.open(BytesIO(response.content))
                    logging.info(f"DEBUG: Successfully created diagram image: {img.size}")
                    return img
                else:
                    error_msg = self._get_error_from_response(response)
                    logging.warning(f"mermaid.ink API error: {error_msg}")
                    
                    # Return the error message so it can be used for AI fixing
                    return {"error": error_msg, "status_code": response.status_code}

            except requests.RequestException as e:
                logging.warning(f"Request failed on attempt {attempt + 1}: {e}")
                if attempt < max_attempts - 1:
                    wait_time = 2 ** attempt  # Exponential backoff
                    logging.info(f"DEBUG: Retrying in {wait_time} seconds...")
                    time.sleep(wait_time)

        return None
    
    def _generate_diagram_from_code(self, mermaid_code: str, max_attempts: int = 5, api_endpoint: str = "mermaid.ink") -> PILImage.Image:
        """
        Generate diagram from mermaid code with retries using specified API.
        
        Args:
            mermaid_code: The mermaid code to render
            max_attempts: Maximum number of attempts
            api_endpoint: Which API to use ("mermaid.ink", "kroki", "quickchart")
            
        Returns:
            PIL Image object or error dict if failed
        """
        # Route to appropriate renderer
        if api_endpoint == "kroki":
            return self._render_with_kroki(mermaid_code)
        elif api_endpoint == "mermaid-cli":
            return self._render_with_mermaid_cli(mermaid_code)
        else: 
            return self._render_with_mermaid_ink(mermaid_code, max_attempts)
        
    def _fix_dot_code_with_ai(self, original_dot_code: str, error_message: str, original_mermaid: str, deployment_name: str = None) -> str:
        """
        Use Azure OpenAI to fix DOT code syntax errors.
        
        Args:
            original_dot_code: The DOT code that failed
            error_message: The error message from Graphviz
            original_mermaid: The original Mermaid code for context
            deployment_name: The Azure OpenAI deployment name to use
            
        Returns:
            Fixed DOT code or original code if AI fixing fails
        """
        try:
            logging.info("DEBUG: Attempting to fix DOT code using Azure OpenAI")
            
            system_prompt = """You are an expert in Graphviz DOT syntax. Your task is to fix invalid DOT code.
            
Rules:
1. Only return the corrected DOT code, nothing else
2. Ensure proper DOT syntax according to Graphviz specifications
3. Maintain the original structure and flow as much as possible
4. Use proper node identifiers (no spaces, use quotes for labels)
5. Ensure proper edge syntax (A -> B;)
6. Fix any syntax errors in the digraph definition
7. Make sure all statements end with semicolons
8. Use proper DOT attributes syntax

Common DOT syntax rules:
- digraph name { ... }
- Node: A [label="text", shape=box];
- Edge: A -> B;
- All statements end with semicolons
- Node IDs should be alphanumeric (use quotes for labels with spaces)
- Valid shapes: box, circle, diamond, ellipse, parallelogram, etc.

Example of valid DOT:
digraph G {
    rankdir=TB;
    A [label="Start", shape=box];
    B [label="Decision", shape=diamond];
    C [label="End", shape=box];
    A -> B;
    B -> C;
}"""

            user_prompt = f"""Fix this DOT code that's causing a Graphviz error:

Original Mermaid Code (for context):
```
{original_mermaid}
```

Current DOT Code:
```
{original_dot_code}
```

Error Message:
{error_message}

Please provide only the corrected DOT code that will work with Graphviz."""

            model_name = deployment_name or get_chat_model_name()
            response = self.azure_client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=1500
            )
            
            # Log usage
            usage = response.usage
            if usage:
                log_openai_usage(
                    module_name=__name__,
                    model=model_name,
                    prompt_tokens=usage.prompt_tokens,
                    completion_tokens=usage.completion_tokens,
                    function_name="_fix_dot_code_with_ai"
                )
            
            fixed_code = response.choices[0].message.content.strip()
            
            # Clean up the response - remove markdown code blocks if present
            if fixed_code.startswith("```"):
                lines = fixed_code.split('\n')
                if lines[0].startswith("```") and lines[-1].strip() == "```":
                    fixed_code = '\n'.join(lines[1:-1])
                elif lines[0].startswith("```"):
                    fixed_code = '\n'.join(lines[1:])
            
            logging.info(f"DEBUG: AI fixed DOT code: {fixed_code[:100]}...")
            return fixed_code.strip()
            
        except Exception as e:
            logging.error(f"Failed to fix DOT code with AI: {e}")
            return original_dot_code
    
    def _fix_mermaid_code_with_ai(self, original_code: str, error_message: str, deployment_name: str = None) -> str:
        """
        Use Azure OpenAI to fix mermaid code syntax errors.
        
        Args:
            original_code: The original mermaid code that failed
            error_message: The error message from the mermaid.ink API
            deployment_name: The Azure OpenAI deployment name to use
            
        Returns:
            Fixed mermaid code or original code if AI fixing fails
        """
        try:
            logging.info("DEBUG: Attempting to fix mermaid code using Azure OpenAI")
            
            system_prompt = """You are an expert in Mermaid diagram syntax. Your task is to fix invalid Mermaid code.
            
Rules:
1. Only return the corrected Mermaid code, nothing else
2. Ensure proper syntax according to Mermaid specifications
3. Maintain the original intent and structure as much as possible
4. Use proper node IDs (alphanumeric, no spaces or special characters except underscores)
5. Ensure proper arrow syntax (-->, -.-> etc.)
6. Fix any syntax errors in the flowchart definition

Common issues to fix:
- Invalid node IDs with spaces or special characters
- Incorrect arrow syntax
- Missing or incorrect flowchart declaration
- Improper quotation marks around labels
- Invalid subgraph syntax"""

            user_prompt = f"""Fix this Mermaid code that's causing an error:

Original Code:
```
{original_code}
```

Error Message:
{error_message}

Please provide only the corrected Mermaid code."""

            model_name = deployment_name or get_chat_model_name()
            response = self.azure_client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=1000
            )
            
            # Log usage
            usage = response.usage
            if usage:
                log_openai_usage(
                    module_name=__name__,
                    model=model_name,
                    prompt_tokens=usage.prompt_tokens,
                    completion_tokens=usage.completion_tokens,
                    function_name="_fix_mermaid_code_with_ai"
                )
            
            fixed_code = response.choices[0].message.content.strip()
            
            # Clean up the response - remove markdown code blocks if present
            if fixed_code.startswith("```"):
                lines = fixed_code.split('\n')
                # Remove first and last lines if they're markdown markers
                if lines[0].startswith("```") and lines[-1].strip() == "```":
                    fixed_code = '\n'.join(lines[1:-1])
                elif lines[0].startswith("```"):
                    fixed_code = '\n'.join(lines[1:])
            
            logging.info(f"DEBUG: AI suggested fixed code: {fixed_code[:100]}...")
            return fixed_code.strip()
            
        except Exception as e:
            logging.error(f"Failed to fix mermaid code with AI: {e}")
            return original_code
        
    def create_mermaid_diagram(self, mermaid_code: str, deployment_name: str = "gpt-4", enable_ai_fix: bool = True, max_ai_attempts: int = 3, use_fallbacks: bool = True) -> PILImage.Image:
        """
        Create a visual diagram from mermaid code with enhanced fallback options.

        Args:
            mermaid_code: The mermaid flowchart code
            deployment_name: Azure OpenAI deployment name for fixing code
            enable_ai_fix: Whether to attempt AI fixing on errors
            max_ai_attempts: Maximum number of AI fixing attempts (reduced for efficiency)
            use_fallbacks: Whether to try alternative rendering services
            
        Returns:
            PIL Image object of the generated diagram or None if all attempts fail.
        """
        try:
            logging.info(f"DEBUG: Creating mermaid diagram with enhanced fallbacks")
            
            # Enhanced list of rendering methods in order of preference
            # Put Graphviz early since it's more reliable than external APIs
            rendering_methods = [
                ("mermaid.ink", "Mermaid.ink API"),
                ("graphviz", "Graphviz (via DOT conversion)"),
                ("mermaid-cli", "Mermaid CLI"),
                ("kroki", "Kroki.io API"),
                ("plantuml", "PlantUML (as last resort)"),
            ]
            
            # If fallbacks are disabled, only use the primary method
            if not use_fallbacks:
                rendering_methods = rendering_methods[:1]
            
            current_code = mermaid_code
            current_dot_code = None  # Track DOT code for Graphviz
            last_error = None
            
            # Try each rendering method
            for method_index, (method_name, method_desc) in enumerate(rendering_methods):
                logging.info(f"DEBUG: Trying rendering method {method_index + 1}/{len(rendering_methods)}: {method_desc}")
                
                # Reset for each method
                method_mermaid_code = current_code if method_index <= 2 else mermaid_code  # Keep improvements for similar methods
                method_dot_code = None
                
                # Try AI fixes for this method
                for ai_attempt in range(max_ai_attempts + 1):
                    if ai_attempt == 0:
                        logging.info(f"DEBUG: {method_desc} - Attempting with current code...")
                    else:
                        logging.info(f"DEBUG: {method_desc} - AI fixing attempt {ai_attempt}/{max_ai_attempts}")
                    
                    # Route to appropriate renderer
                    if method_name == "graphviz":
                        result = self._render_with_graphviz(
                            method_mermaid_code, 
                            dot_code=method_dot_code,
                            attempt_num=ai_attempt + 1
                        )
                    elif method_name == "plantuml":
                        result = self._render_with_plantuml_fallback(method_mermaid_code)
                    else:
                        result = self._generate_diagram_from_code(method_mermaid_code, api_endpoint=method_name)
                    
                    # If successful, return the image
                    if isinstance(result, PILImage.Image):
                        success_msg = f"✓ Success with {method_desc}"
                        if method_index > 0:
                            success_msg += f" (fallback #{method_index + 1})"
                        if ai_attempt > 0:
                            success_msg += f" after {ai_attempt} AI fix(es)"
                        logging.info(success_msg)
                        return result
                    
                    # Handle failures and AI fixing
                    if isinstance(result, dict) and "error" in result:
                        last_error = result["error"]
                        
                        if ai_attempt >= max_ai_attempts or not enable_ai_fix:
                            break
                        
                        logging.info(f"DEBUG: {method_desc} failed, trying AI fix...")
                        
                        # Use method-specific AI fixing
                        if method_name == "graphviz" and "dot_code" in result:
                            # Fix the DOT code specifically
                            method_dot_code = self._fix_dot_code_with_ai(
                                original_dot_code=result["dot_code"],
                                error_message=last_error,
                                original_mermaid=result.get("original_mermaid", method_mermaid_code),
                                deployment_name=deployment_name
                            )
                            logging.info("DEBUG: AI provided fixed DOT code")
                        else:
                            # Fix the Mermaid code
                            method_mermaid_code = self._fix_mermaid_code_with_ai(
                                original_code=method_mermaid_code,
                                error_message=last_error,
                                deployment_name=deployment_name
                            )
                        
                        time.sleep(0.5)  # Brief pause between attempts
                    else:
                        # Method failed or not available
                        if result is None:
                            logging.info(f"DEBUG: {method_desc} not available or failed")
                        break
            
            # If all methods failed
            logging.error("All rendering methods failed")
            return None
            
        except Exception as e:
            logging.error(f"Error creating mermaid diagram: {e}")
            import traceback
            logging.error(f"Full error: {traceback.format_exc()}")
            return None
        
class MermaidDiagramGenerator_v1:
    def __init__(self):
        """
        Initialize the Mermaid Diagram Generator with Azure OpenAI credentials.
        
        Args:
            azure_endpoint: Azure OpenAI endpoint URL
            azure_api_key: Azure OpenAI API key
            azure_api_version: Azure OpenAI API version
        """

        self.azure_client = get_openai_client()
        self.fallback_apis = [
                "https://mermaid.ink/img/",
                "https://kroki.io/mermaid/svg/",
                "https://quickchart.io/mermaid?chart="
            ]
    
    def _render_with_kroki(self, mermaid_code: str) -> PILImage.Image:
        """
        Render mermaid diagram using Kroki.io API via POST request.
        
        Args:
            mermaid_code: The mermaid code to render
            
        Returns:
            PIL Image object or None if failed
        """
        try:
            logging.info("DEBUG: Attempting to render with Kroki.io via POST")

            api_url = 'https://kroki.io/mermaid/png'
            headers = {'Content-Type': 'text/plain'}

            response = requests.post(api_url, data=mermaid_code.encode('utf-8'), headers=headers, timeout=15)

            if response.status_code == 200:
                img = PILImage.open(BytesIO(response.content))
                logging.info(f"DEBUG: Kroki.io success: {img.size}")
                return img
            else:
                logging.warning(f"Kroki.io failed: {response.status_code} - {response.text}")
                return None

        except Exception as e:
            logging.error(f"Kroki.io error: {e}")
            return None

    def _render_with_mermaid_cli(self, mermaid_code: str) -> PILImage.Image:
        """
        Render mermaid diagram using local Mermaid CLI (if installed).
        
        Args:
            mermaid_code: The mermaid code to render
            
        Returns:
            PIL Image object or None if failed
        """
        try:
            logging.info("DEBUG: Attempting to render with Mermaid CLI")
            
            # Check if mermaid CLI is available
            # try:
            #     subprocess.run(['npx', 'mmdc', '--version'], capture_output=True, check=True)
            # except (subprocess.CalledProcessError, FileNotFoundError):
            #     logging.info("DEBUG: Mermaid CLI not available, skipping")
            #     return None
            
            # Create temporary files
            with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as input_file:
                input_file.write(mermaid_code)
                input_path = input_file.name
            
            output_path = input_path.replace('.mmd', '.png')
            
            try:
                # Run mermaid CLI
                result = subprocess.run([
                    'npx','mmdc', 
                    '-i', input_path, 
                    '-o', output_path,
                    '-t', 'neutral',
                    '-b', 'white'
                ], capture_output=True, text=True, timeout=30)
#                 result = subprocess.run(
#     f'npx mmdc -i "{input_path}" -o "{output_path}" -t neutral -b white',
#     capture_output=True,
#     text=True,
#     timeout=30,
#     shell=True  # Important for Windows
# )

                
                if result.returncode == 0 and os.path.exists(output_path):
                    img = PILImage.open(output_path)
                    logging.info(f"DEBUG: Mermaid CLI success: {img.size}")
                    return img.copy()  # Copy to memory before file cleanup
                else:
                    logging.warning(f"Mermaid CLI failed: {result.stderr}")
                    return None
                    
            finally:
                # Cleanup temporary files
                for path in [input_path, output_path]:
                    if os.path.exists(path):
                        os.unlink(path)
                        
        except Exception as e:
            logging.error(f"Mermaid CLI error: {e}")
            return None
    
    def _fix_mermaid_code_with_ai(self, original_code: str, error_message: str, deployment_name: str = "gpt-4o") -> str:
        """
        Use Azure OpenAI to fix mermaid code syntax errors.
        
        Args:
            original_code: The original mermaid code that failed
            error_message: The error message from the mermaid.ink API
            deployment_name: The Azure OpenAI deployment name to use
            
        Returns:
            Fixed mermaid code or original code if AI fixing fails
        """
        try:
            logging.info("DEBUG: Attempting to fix mermaid code using Azure OpenAI")
            
            system_prompt = """You are an expert in Mermaid diagram syntax. Your task is to fix invalid Mermaid code.
            
Rules:
1. Only return the corrected Mermaid code, nothing else
2. Ensure proper syntax according to Mermaid specifications
3. Maintain the original intent and structure as much as possible
4. Use proper node IDs (alphanumeric, no spaces or special characters except underscores)
5. Ensure proper arrow syntax (-->, -.-> etc.)
6. Fix any syntax errors in the flowchart definition

Common issues to fix:
- Invalid node IDs with spaces or special characters
- Incorrect arrow syntax
- Missing or incorrect flowchart declaration
- Improper quotation marks around labels
- Invalid subgraph syntax"""

            user_prompt = f"""Fix this Mermaid code that's causing an error:

Original Code:
```
{original_code}
```

Error Message:
{error_message}

Please provide only the corrected Mermaid code."""

            response = self.azure_client.chat.completions.create(
                model=deployment_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=1000
            )
            usage = response.usage
            if usage:
                log_openai_usage(
                    module_name=__name__,
                    model=deployment_name,
                    prompt_tokens=usage.prompt_tokens,
                    completion_tokens=usage.completion_tokens,
                    function_name="_fix_mermaid_code_with_ai"
                )
            
            fixed_code = response.choices[0].message.content.strip()
            
            # Clean up the response - remove markdown code blocks if present
            if fixed_code.startswith("```"):
                lines = fixed_code.split('\n')
                # Remove first and last lines if they're markdown markers
                if lines[0].startswith("```") and lines[-1].strip() == "```":
                    fixed_code = '\n'.join(lines[1:-1])
                elif lines[0].startswith("```"):
                    fixed_code = '\n'.join(lines[1:])
            
            logging.info(f"DEBUG: AI suggested fixed code: {fixed_code[:100]}...")
            return fixed_code.strip()
            
        except Exception as e:
            logging.error(f"Failed to fix mermaid code with AI: {e}")
            return original_code
    
    def _get_error_from_response(self, response) -> str:
        """
        Extract error information from the mermaid.ink API response.
        
        Args:
            response: The HTTP response from mermaid.ink
            
        Returns:
            Error message string
        """
        try:
            if response.status_code == 400:
                return f"Bad Request (400): Invalid mermaid syntax. Response: {response.text[:200]}"
            elif response.status_code == 500:
                return f"Server Error (500): Mermaid processing failed. Response: {response.text[:200]}"
            else:
                return f"HTTP {response.status_code}: {response.text[:200]}"
        except:
            return f"HTTP {response.status_code}: Unable to parse error response"
    
    def _generate_diagram_from_code(self, mermaid_code: str, max_attempts: int = 5, api_endpoint: str = "mermaid.ink") -> PILImage.Image:
        """
        Generate diagram from mermaid code with retries using specified API.
        
        Args:
            mermaid_code: The mermaid code to render
            max_attempts: Maximum number of attempts
            api_endpoint: Which API to use ("mermaid.ink", "kroki", "quickchart")
            
        Returns:
            PIL Image object or error dict if failed
        """
        # Route to appropriate renderer
        if api_endpoint == "kroki":
            return self._render_with_kroki(mermaid_code)
        elif api_endpoint == "mermaid-cli":
            return self._render_with_mermaid_cli(mermaid_code)
        else:
            # Default to mermaid.ink with retry logic
            return self._render_with_mermaid_ink(mermaid_code, max_attempts)

    
    def _render_with_mermaid_ink(self, mermaid_code: str, max_attempts: int = 5) -> PILImage.Image:
        """
        Generate diagram from mermaid code using mermaid.ink with retries.
        
        Args:
            mermaid_code: The mermaid code to render
            max_attempts: Maximum number of attempts
            
        Returns:
            PIL Image object or error dict if failed
        """
        # Clean up the mermaid code
        mermaid_code = mermaid_code.strip()
        
        # Encode the mermaid code for the API
        graphbytes = mermaid_code.encode("utf8")
        base64_bytes = base64.urlsafe_b64encode(graphbytes)
        base64_string = base64_bytes.decode("ascii")
        
        # Get the diagram image from mermaid.ink API
        api_url = f'https://mermaid.ink/img/{base64_string}'

        for attempt in range(max_attempts):
            logging.info(f"DEBUG: Mermaid.ink attempt {attempt + 1} - Calling API: {api_url[:100]}...")
            try:
                response = requests.get(api_url, timeout=15)
                logging.info(f"DEBUG: mermaid.ink API response status: {response.status_code}")

                if response.status_code == 200:
                    img = PILImage.open(BytesIO(response.content))
                    logging.info(f"DEBUG: Successfully created diagram image: {img.size}")
                    return img
                else:
                    error_msg = self._get_error_from_response(response)
                    logging.warning(f"mermaid.ink API error: {error_msg}")
                    
                    # Return the error message so it can be used for AI fixing
                    return {"error": error_msg, "status_code": response.status_code}

            except requests.RequestException as e:
                logging.warning(f"Request failed on attempt {attempt + 1}: {e}")
                if attempt < max_attempts - 1:
                    wait_time = 2 ** attempt  # Exponential backoff
                    logging.info(f"DEBUG: Retrying in {wait_time} seconds...")
                    time.sleep(wait_time)

        return None
    
    def create_mermaid_diagram(self, mermaid_code: str, deployment_name: str = "gpt-4o", enable_ai_fix: bool = True, max_ai_attempts: int = 5, use_fallbacks: bool = True) -> PILImage.Image:
        """
        Create a visual diagram from mermaid code with AI-powered error correction and multiple fallback renderers.

        Args:
            mermaid_code: The mermaid flowchart code
            deployment_name: Azure OpenAI deployment name for fixing code
            enable_ai_fix: Whether to attempt AI fixing on errors
            max_ai_attempts: Maximum number of AI fixing attempts
            use_fallbacks: Whether to try alternative rendering services
            
        Returns:
            PIL Image object of the generated diagram or None if all attempts fail.
        """
        try:
            logging.info(f"DEBUG: Creating mermaid diagram with code: {mermaid_code}")
            
            # List of rendering methods to try in order
            rendering_methods = [
                ("mermaid.ink", "Mermaid.ink API"),
                ("kroki", "Kroki.io API"),
                ("mermaid-cli", "Mermaid CLI"),
            ]
            
            # If fallbacks are disabled, only use the primary method
            if not use_fallbacks:
                rendering_methods = rendering_methods[:1]
            
            current_code = mermaid_code
            last_error = None
            
            # Try each rendering method
            for method_index, (method_name, method_desc) in enumerate(rendering_methods):
                logging.info(f"DEBUG: Trying rendering method: {method_desc}")
                
                # For each rendering method, try AI fixes if enabled
                for ai_attempt in range(max_ai_attempts + 1):  # +1 for original attempt
                    if ai_attempt == 0:
                        logging.info(f"DEBUG: {method_desc} - Attempting with {'original' if method_index == 0 else 'current'} code...")
                    else:
                        logging.info(f"DEBUG: {method_desc} - AI fixing attempt {ai_attempt}/{max_ai_attempts}")
                    
                    # Try to generate diagram with current code and method
                    result = self._generate_diagram_from_code(current_code, api_endpoint=method_name)
                    
                    # If successful, return the image
                    if isinstance(result, PILImage.Image):
                        success_msg = f"Success with {method_desc}"
                        if method_index > 0:
                            success_msg += f" (fallback method {method_index + 1})"
                        if ai_attempt > 0:
                            success_msg += f" after {ai_attempt} AI fix attempt(s)"
                        logging.info(f"DEBUG: {success_msg}!")
                        return result
                    
                    # If it failed and we have more AI attempts for this method, try AI fixing
                    if isinstance(result, dict) and "error" in result:
                        last_error = result["error"]
                        
                        # If this was our last AI attempt for this method or AI fixing is disabled, break to try next method
                        if ai_attempt >= max_ai_attempts or not enable_ai_fix:
                            break
                        
                        logging.info(f"DEBUG: {method_desc} attempt {ai_attempt + 1} failed, trying AI fix...")
                        
                        # Get AI to fix the code
                        fixed_code = self._fix_mermaid_code_with_ai(
                            original_code=current_code,
                            error_message=last_error,
                            deployment_name=deployment_name
                        )
                        
                        # If AI provided the same code, no point in retrying with this method
                        # if fixed_code == current_code:
                        #     logging.info(f"DEBUG: AI returned same code for {method_desc}, trying next method")
                        #     break
                        
                        current_code = fixed_code
                        logging.info(f"DEBUG: AI provided new code for {method_desc} attempt {ai_attempt + 2}")
                        
                        # Add a small delay between AI attempts
                        if ai_attempt < max_ai_attempts:
                            time.sleep(1)
                    else:
                        # Method is not available or other issue, try next method
                        logging.info(f"DEBUG: {method_desc} not available, trying next method")
                        break
                
                # Reset to original code for next rendering method (each method gets a fresh start)
                if method_index == 0:  # After first method, keep any AI improvements
                    pass  # Keep the AI-improved code for other methods
                
            logging.error(f"DEBUG: All rendering methods failed. Last error: {last_error}")
            return None

        except Exception as e:
            logging.error(f"Error creating mermaid diagram: {e}")
            import traceback
            logging.error(f"Full error: {traceback.format_exc()}")
            return None

                
class DocumentGenerator:
    """Class for generating product documentation from screenshots and video content"""
    
    def __init__(self, video_path: str, screenshots: List[Tuple[Any, float, str]],
                 use_ai: bool = False, title: str = "", description: str = "",
                 speech_segments: List[Tuple[float, str]] = [], document_type: str = "kt_document",
                 generate_missing_questions: bool = False, generate_process_map: bool = False,
                 include_screenshots: bool = False, meeting_participants: Optional[List[str]] = None,
                 meeting_highlights: Optional[List[str]] = None,
                 meeting_duration_minutes: Optional[float] = None,
                 session_guid: Optional[str] = None, teams_llm_config: Optional[Dict[str, Any]] = None):
        """
        Initialize the document generator
        
        Args:
            video_path: Path to the video file
            screenshots: List of tuples (image, timestamp, reason)
            use_ai: Whether to use AI for generating descriptions
            title: Optional title for the document
            description: Optional description for the document
            speech_segments: Optional list of tuples (timestamp, speech_text) for narrative documentation
            document_type: Type of document to generate ("kt_document" or "meeting_summary")
            generate_missing_questions: Whether to generate missing questions section
            generate_process_map: Whether to generate process map diagram
            teams_llm_config: LLM configurations of the team
        """
        self.video_path = video_path
        self.screenshots = screenshots
        self.title = title if title else os.path.basename(video_path).split('.')[0]
        self.description = description if description else f"Documentation generated from {os.path.basename(video_path)}"
        self.speech_segments = speech_segments if speech_segments else []
        self.document_type = document_type
        self.generate_missing_questions = generate_missing_questions
        self.generate_process_map = generate_process_map
        self.include_screenshots = include_screenshots
        self.meeting_participants = meeting_participants or []
        self.meeting_highlights = meeting_highlights or []
        self.meeting_duration_minutes = meeting_duration_minutes or 0
        self.session_guid = session_guid
        self.teams_llm_config = teams_llm_config or {}
        
        # The screenshots should already be deduplicated in app.py
        # Just make sure they're sorted by timestamp
        self.screenshots.sort(key=lambda x: x[1])
        
        # Initialize OpenAI client if available
        self.model = get_chat_model_name()
        self.completion = completion

        self._litellm_model = teams_llm_config.get("model")
        self.auth_token = self.teams_llm_config.pop("auth_token", None)
        self._litellm_kwargs= self.teams_llm_config

        # if self.model:
        #     self._litellm_model = f"azure/{self.model}" if USE_AZURE else self.model

        # if USE_AZURE and AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT:
        #     self._litellm_kwargs = {
        #         "api_key": AZURE_OPENAI_API_KEY,
        #         "api_base": AZURE_OPENAI_ENDPOINT,
        #         "base_url": AZURE_OPENAI_ENDPOINT,
        #         "api_version": AZURE_OPENAI_API_VERSION or "2024-02-01"
        #     }
        # elif not USE_AZURE and OPENAI_API_KEY:
        #     self._litellm_kwargs = {"api_key": OPENAI_API_KEY}

        self.use_ai = bool(use_ai and self._litellm_model and self._litellm_kwargs)


    def _invoke_completion(self, messages: List[Dict[str, Any]], **kwargs):
        """
        Wrapper around litellm.completion that injects the correct model and Azure/OpenAI credentials.
        """
        # if not self.use_ai or not self._litellm_model:
        #     raise RuntimeError("AI completion requested but configuration is unavailable.")
        
        request_kwargs = dict(self._litellm_kwargs)
        request_kwargs.update(kwargs)
        # request_kwargs["model"] = self._litellm_model
        request_kwargs["messages"] = messages
        return self.completion(**request_kwargs)
    
    def _clean_json_response(self, content: str) -> str:
        """
        Remove Markdown code fences (``` or ```json) that models sometimes wrap around JSON responses.
        """
        if not content:
            return content
        
        cleaned = content.strip()
        if cleaned.startswith("```"):
            lines = cleaned.splitlines()
            # Drop the opening fence line (e.g., ```json)
            lines = lines[1:]
            # Remove any trailing fence markers
            while lines and lines[-1].strip() == "```":
                lines.pop()
            cleaned = "\n".join(lines).strip()
        return cleaned
    

    def generate_mermaid_editor_url_docx(self, mermaid_code):
        """
        Generate a Mermaid Chart editor URL with the provided mermaid code
        """
        try:
            # Encode the mermaid code for URL
            encoded_code = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
            # Create the editor URL with the encoded diagram
            editor_url = f"https://mermaid.live/edit#base64:{encoded_code}"
            return editor_url
        except Exception as e:
            logging.warning(f"Failed to generate mermaid editor URL: {e}")
            # Fallback to basic editor
            return "https://mermaid.live/edit"

    def add_hyperlink(self, paragraph, text, url):
        """
        Add a hyperlink to a paragraph in Word document
        """
        try:
            # Get the paragraph's hyperlink collection
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            
            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Create run element for the hyperlink text
            run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Style the hyperlink (blue and underlined)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
            run.append(rPr)
            run.text = text
            hyperlink.append(run)
            
            paragraph._p.append(hyperlink)
            return True
        except Exception as e:
            logging.warning(f"Failed to add hyperlink: {e}")
            # Fallback: add as plain text
            paragraph.add_run(text)
            return False

    def generate_mermaid_editor_url(self, mermaid_code):
        """
        Generate a Mermaid Chart editor URL with the provided mermaid code
        """
        try:
            # Encode the mermaid code for URL
            encoded_code = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
            # Create the editor URL with the encoded diagram
            editor_url = f"https://mermaid.live/edit#pako:{encoded_code}"
            return editor_url
        except Exception as e:
            logging.warning(f"Failed to generate mermaid editor URL: {e}")
            # Fallback to basic editor
            return "https://mermaid.live/edit"
    
    def _calculate_optimal_image_size(self, img: PILImage.Image, max_width_inches: float = 7.0, max_height_inches: float = 5.0) -> tuple:
            """
            Calculate optimal image dimensions while maintaining aspect ratio
            
            Args:
                img: PIL Image object
                max_width_inches: Maximum width in inches
                max_height_inches: Maximum height in inches
                
            Returns:
                Tuple of (width_inches, height_inches)
            """
            # Get original dimensions
            original_width, original_height = img.size
            aspect_ratio = original_width / original_height
            
            # Calculate dimensions based on max constraints
            if aspect_ratio > 1:  # Wider than tall
                width_inches = min(max_width_inches, original_width / 100)  # Assume 100 DPI
                height_inches = width_inches / aspect_ratio
            else:  # Taller than wide
                height_inches = min(max_height_inches, original_height / 100)  # Assume 100 DPI
                width_inches = height_inches * aspect_ratio
            
            # Ensure minimum readable size
            min_width, min_height = 2.0, 1.5
            width_inches = max(min_width, width_inches)
            height_inches = max(min_height, height_inches)
            
            # Ensure we don't exceed maximum constraints
            if width_inches > max_width_inches:
                width_inches = max_width_inches
                height_inches = width_inches / aspect_ratio
            
            if height_inches > max_height_inches:
                height_inches = max_height_inches
                width_inches = height_inches * aspect_ratio
            
            return width_inches, height_inches

    def _create_mermaid_diagram(self, mermaid_code: str) -> PILImage.Image:
        """
        Create a visual diagram from mermaid code using mermaid.ink API with retries.

        Args:
            mermaid_code: The mermaid flowchart code
            
        Returns:
            PIL Image object of the generated diagram or None if all attempts fail.
        """
        try:
            logging.info(f"DEBUG: Creating mermaid diagram with code: {mermaid_code}")
            
            # Clean up the mermaid code
            mermaid_code = mermaid_code.strip()
            
            # Encode the mermaid code for the API
            graphbytes = mermaid_code.encode("utf8")
            base64_bytes = base64.urlsafe_b64encode(graphbytes)
            base64_string = base64_bytes.decode("ascii")
            
            # Get the diagram image from mermaid.ink API
            api_url = f'https://mermaid.ink/img/{base64_string}'

            attempts = 3
            for attempt in range(attempts):
                logging.info(f"DEBUG: Attempt {attempt + 1} - Calling mermaid.ink API: {api_url[:100]}...")
                try:
                    response = requests.get(api_url, timeout=10)
                    logging.info(f"DEBUG: mermaid.ink API response status: {response.status_code}")

                    if response.status_code == 200:
                        img = PILImage.open(BytesIO(response.content))
                        logging.info(f"DEBUG: Successfully created diagram image: {img.size}")
                        return img
                    else:
                        logging.info(f"Error from mermaid.ink API: Status code {response.status_code}")
                        logging.info(f"Response content: {response.text[:200]}...")

                except requests.RequestException as e:
                    logging.info(f"Request failed: {e}")

                if attempt < attempts - 1:
                    wait_time = 2 ** attempt  # Exponential backoff (1s, 2s)
                    logging.info(f"DEBUG: Retrying in {wait_time} seconds...")
                    time.sleep(wait_time)

            logging.info("DEBUG: All attempts failed. Returning None.")
            return None

        except Exception as e:
            logging.info(f"Error creating mermaid diagram: {e}")
            import traceback
            logging.info(f"Full error: {traceback.format_exc()}")
            return None

    def _calculate_optimal_image_size_pdf(self, img: PILImage.Image, max_width_inches: float = 7.0, max_height_inches: float = 5.0, min_width_inches: float = 4.0, min_height_inches: float = 3.0) -> tuple:
            """
            Calculate optimal image dimensions while maintaining aspect ratio and ensuring readability.

            Args:
                img: PIL Image object
                max_width_inches: Maximum width in inches
                max_height_inches: Maximum height in inches
                min_width_inches: Minimum width in inches for readability
                min_height_inches: Minimum height in inches for readability

            Returns:
                Tuple of (width_inches, height_inches)
            """
            # Get original dimensions in pixels
            original_width, original_height = img.size
            if original_height == 0:  # Prevent division by zero
                original_height = 1
            aspect_ratio = original_width / original_height

            # Use DPI from image metadata if available, otherwise default to 150 for better text clarity
            dpi = img.info.get('dpi', (150, 150))[0] if hasattr(img, 'info') and 'dpi' in img.info else 150
            max_width_pixels = max_width_inches * dpi
            max_height_pixels = max_height_inches * dpi
            min_width_pixels = min_width_inches * dpi
            min_height_pixels = min_height_inches * dpi

            # Scale to fit within max dimensions while preserving aspect ratio
            if aspect_ratio > (max_width_pixels / max_height_pixels):
                # Image is wider relative to max dimensions, scale by width
                width_pixels = max_width_pixels
                height_pixels = width_pixels / aspect_ratio
            else:
                # Image is taller relative to max dimensions, scale by height
                height_pixels = max_height_pixels
                width_pixels = height_pixels * aspect_ratio

            # Convert to inches
            width_inches = width_pixels / dpi
            height_inches = height_pixels / dpi

            # Ensure minimum size for readability
            if width_inches < min_width_inches:
                width_inches = min_width_inches
                height_inches = width_inches / aspect_ratio
            if height_inches < min_height_inches:
                height_inches = min_height_inches
                width_inches = height_inches * aspect_ratio

            # Re-check max constraints to avoid exceeding page limits
            if width_inches > max_width_inches:
                width_inches = max_width_inches
                height_inches = width_inches / aspect_ratio
            if height_inches > max_height_inches:
                height_inches = max_height_inches
                width_inches = height_inches * aspect_ratio

            # Log dimensions for debugging
            logging.info(f"Original: {img.size}, Aspect Ratio: {aspect_ratio:.2f}, DPI: {dpi}")
            logging.info(f"Calculated: ({width_inches:.2f}in, {height_inches:.2f}in)")

            return width_inches, height_inches

    def _format_content_for_pdf(self, content: str, normal_style, bold_style) -> List:
        """
        Format content with bold headings and bullet points for PDF
        
        Args:
            content: Content string with markdown-style formatting
            normal_style: Normal paragraph style
            bold_style: Bold paragraph style
            
        Returns:
            List of formatted paragraph elements
        """
        elements = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 0.1*inch))
                continue
                
            # Handle bold headings (**TEXT**)
            if line.startswith('**') and line.endswith('**'):
                heading_text = line.strip('*')
                elements.append(Paragraph(heading_text, bold_style))
                elements.append(Spacer(1, 0.1*inch))
            # Handle bullet points (• TEXT)
            elif line.startswith('• ') or line.startswith('- '):
                bullet_text = line[2:]
                elements.append(Paragraph(f"• {bullet_text}", normal_style))
            # Handle numbered items
            elif line and line[0].isdigit() and '. ' in line:
                elements.append(Paragraph(line, normal_style))
            # Regular paragraphs
            else:
                elements.append(Paragraph(line, normal_style))
        
        return elements

    def _extract_keywords_from_reason(self, reason):
        """Extract meaningful keywords from screenshot reason for contextual matching"""
        if not reason:
            return []
        
        # Common keywords that indicate specific UI elements or actions
        ui_keywords = ['login', 'dashboard', 'menu', 'button', 'form', 'page', 'screen', 'navigation', 'click', 'input', 'search', 'settings', 'profile', 'home', 'report', 'data', 'table', 'chart', 'graph', 'filter', 'export', 'import', 'save', 'edit', 'delete', 'create', 'add', 'view', 'list', 'detail', 'summary']
        
        keywords = []
        reason_words = reason.lower().split()
        
        # Extract UI-related keywords
        for word in reason_words:
            # Remove punctuation and get clean word
            clean_word = ''.join(char for char in word if char.isalnum())
            if len(clean_word) > 3 and clean_word in ui_keywords:
                keywords.append(clean_word)
        
        # Add specific patterns
        if 'keyword trigger' in reason.lower():
            # Extract the actual triggered keyword
            if ':' in reason:
                triggered_text = reason.split(':', 1)[1].strip()
                keywords.extend(triggered_text.lower().split()[:3])  # First 3 words
        
        return keywords[:5]  # Return top 5 keywords

    def _enhance_screenshot_reasons(self, speech_segments):
        """
        Enhance screenshot reasons using AI analysis for better contextual descriptions
        
        Args:
            speech_segments: List of (timestamp, text) tuples
            
        Returns:
            List of enhanced screenshots with better contextual reasons
        """
        if not self.use_ai:
            return self.screenshots
        
        enhanced_screenshots = []
        
        try:
            # Prepare context for AI analysis
            for img, ts, original_reason in self.screenshots:
                # Get speech context around this timestamp (10 seconds before and after)
                context_speech = []
                for speech_ts, speech_text in speech_segments:
                    if abs(speech_ts - ts) <= 10:  # Within 10 seconds
                        context_speech.append(f"[{speech_ts:.1f}s] {speech_text}")
                
                speech_context = "\n".join(context_speech) if context_speech else "No speech detected nearby"
                
                # Determine screenshot type and enhance reason
                enhanced_reason = original_reason
                
                if "sampled frame" in original_reason.lower():
                    # These are generic samples - enhance with AI
                    if self.use_ai and speech_context != "No speech detected nearby":
                        enhanced_reason = self._get_ai_enhanced_reason(ts, speech_context, "visual content analysis")
                    else:
                        enhanced_reason = f"Key visual moment at {ts:.1f}s"
                
                elif "scene change" in original_reason.lower():
                    # Scene change detected - enhance with context
                    if speech_context != "No speech detected nearby":
                        enhanced_reason = self._get_ai_enhanced_reason(ts, speech_context, "scene transition")
                    else:
                        enhanced_reason = f"Scene transition detected - {original_reason}"
                
                elif "keyword trigger" in original_reason.lower():
                    if speech_context != "No speech detected nearby":
                        enhanced_reason = self._get_ai_enhanced_reason(ts, speech_context, "scene transition")
                    else:
                        enhanced_reason = original_reason.replace("Keyword trigger: ", "Speech keyword detected: ")
                                
                elif "ai detected" in original_reason.lower():
                    # AI detection - already enhanced, keep as is
                    enhanced_reason = original_reason
                
                else:
                    # Other types (mouse clicks, UI changes, etc.) - enhance if we have speech context
                    if speech_context != "No speech detected nearby":
                        enhanced_reason = self._get_ai_enhanced_reason(ts, speech_context, "user interaction")
                    else:
                        enhanced_reason = original_reason
                
                enhanced_screenshots.append((img, ts, enhanced_reason))
                
        except Exception as e:
            logging.info(f"Error enhancing screenshot reasons: {e}")
            return self.screenshots
        
        return enhanced_screenshots

    def _get_ai_enhanced_reason(self, timestamp, speech_context, detection_type):
        """
        Get AI-enhanced reason for a screenshot based on speech context
        
        Args:
            timestamp: Screenshot timestamp
            speech_context: Speech text around the timestamp
            detection_type: Type of detection (visual, scene transition, user interaction)
            
        Returns:
            Enhanced contextual reason string
        """
        try:
            prompt = f"""
            Analyze this screenshot moment and provide a brief, contextual description.
            
            Timestamp: {timestamp:.1f} seconds
            Detection Type: {detection_type}
            Speech Context:
            {speech_context}
            
            Provide a concise reason (max 60 characters) that explains what's likely happening at this moment based on the speech context. Focus on:
            - What action or feature is being demonstrated
            - What UI element or screen is being shown
            - What process or workflow step is occurring
            
            Format: Brief descriptive phrase (no "Screenshot of" or "Image showing")
            Examples: "Login screen demonstration", "Dashboard navigation", "Settings configuration", "Report generation process"
            """
            
            response = self._invoke_completion(
                messages=[
                    {"role": "system", "content": "You are an expert at analyzing video content and providing concise, contextual descriptions."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=50,
                temperature=0.3
            )
            
            # usage = response.usage
            # if usage:
            #     log_openai_usage(
            #         module_name=__name__,
            #         model=get_chat_model_name(),
            #         prompt_tokens=usage.prompt_tokens,
            #         completion_tokens=usage.completion_tokens,
            #         function_name="_get_ai_enhanced_reason",
            #         session_id=self.session_guid
            #     )
            if response:
                token_tracker.track_response(response=response, auth_token=self.auth_token, model= self._litellm_model)
            enhanced_reason = response.choices[0].message.content.strip()
            
            # Ensure it's not too long
            if len(enhanced_reason) > 60:
                enhanced_reason = enhanced_reason[:57] + "..."
            
            return enhanced_reason
            
        except Exception as e:
            logging.info(f"Error getting AI enhanced reason: {e}")
            return f"Key moment at {timestamp:.1f}s ({detection_type})"

    def _add_formatted_text_to_docx(self, doc, text_content):
        """
        Add formatted text to DOCX document with proper markdown handling
        
        Args:
            doc: Document object
            text_content: Text content with markdown formatting
        """
        import re
        
        lines = text_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Handle full line bold headings (**TEXT**)
            if line.startswith('**') and line.endswith('**') and line.count('**') == 2:
                heading_text = line.strip('*')
                para = doc.add_paragraph()
                para.add_run(heading_text).bold = True
            # Handle bullet points
            elif line.startswith('• ') or line.startswith('- '):
                bullet_text = line[2:]
                # Process inline bold in bullet points
                self._add_paragraph_with_inline_formatting(doc, bullet_text, style='List Bullet')
            # Handle numbered items and other content with inline formatting
            else:
                self._add_paragraph_with_inline_formatting(doc, line)

    def _add_paragraph_with_inline_formatting(self, doc, text, style=None):
        """
        Add a paragraph with inline bold formatting support
        
        Args:
            doc: Document object
            text: Text with potential **bold** markdown
            style: Optional paragraph style
        """
        import re
        
        para = doc.add_paragraph(style=style)
        
        # Split text by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # This is bold text
                bold_text = part.strip('*')
                para.add_run(bold_text).bold = True
            elif part:
                # This is regular text
                para.add_run(part)

    def _generate_missing_questions(self) -> str:
        """
        Generate missing questions that should be asked in future meetings based on the transcript.
        
        Returns:
            String containing formatted missing questions
        """
        if not self.use_ai or not self.speech_segments:
            return ""
            
        # Combine all speech segments into a full transcript
        full_transcript = "\n".join([text for _, text in self.speech_segments])
        
        prompt = f"""
        Based on the following meeting transcript, identify important questions that should be asked in future meetings to gather missing information.
        
        Transcript:
        {full_transcript}
        
        Please provide well-formatted questions organized by category. Use this exact format:
        
        **CLARIFICATION QUESTIONS**
        • What specific [topic] details need clarification?
        • How should [process] be implemented exactly?
        
        **TECHNICAL SPECIFICATIONS**
        • What are the exact technical requirements for [feature]?
        • Which technologies or frameworks should be used?
        
        **TIMELINE & RESOURCES**
        • What is the realistic timeline for [deliverable]?
        • Who will be responsible for [task/area]?
        
        **IMPLEMENTATION DETAILS**
        • How should [feature] integrate with existing systems?
        • What are the testing and validation requirements?
        
        **RISKS & MITIGATION**
        • What potential risks should be considered?
        • What backup plans need to be established?
        
        Focus only on actionable questions that would provide valuable missing information.
        """
        
        try:
            response = self._invoke_completion(
                messages=[
                    {"role": "system", "content": "You are an expert meeting analyst who identifies missing questions and areas that need follow-up."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1500,
                temperature=0.3
            )
            # usage = response.usage
            # if usage:
            #     log_openai_usage(
            #         module_name=__name__,
            #         model=get_chat_model_name(),
            #         prompt_tokens=usage.prompt_tokens,
            #         completion_tokens=usage.completion_tokens,
            #         function_name="_generate_missing_questions",
            #         session_id=self.session_guid
            #     )

            if response:
                token_tracker.track_response(response=response, auth_token=self.auth_token, model=self._litellm_model)
            
            return response.choices[0].message.content
        except Exception as e:
            logging.info(f"Error generating missing questions: {e}")
            return "## Missing Questions for Next Meeting\n\nError generating questions. Please check OpenAI configuration."

    def _generate_process_map(self) -> str:
        """
        Generate a clean process map diagram if a process is explained in the video.
        
        Returns:
            String containing text-based process map with mermaid diagram
        """
        if not self.use_ai or not self.speech_segments:
            return ""
            
        # Combine all speech segments into a full transcript
        full_transcript = "\n".join([text for _, text in self.speech_segments])
        
        mermaid_template = """
        
        **PROCESS DETAILS**
        
        Process Name : [Name of the identified process]
        Key Stakeholders: [People/roles involved]
        Prerequisites: [What's needed to start]
        Expected Outcome: [What the process produces]
        Estimated Duration: [Time to complete]
        
        **Step Descriptions:**
        1. Step 1: Detailed explanation
        2. Step 2: Detailed explanation
        3. Step 3: Detailed explanation
        
        Use proper mermaid syntax with meaningful node IDs and clear flow connections. If no clear process is identified, respond with "No process workflow identified in the transcript."
        """
        
        prompt = f"""
        Analyze the following transcript and create a mermaid flowchart diagram if any business process or workflow is explained. Create mermaid code without any syntax error. Be doubly careful with the brackets in the code as it causes errors.
        
        Transcript:
        {full_transcript}
        
        If you identify a process or workflow, create both a mermaid diagram and description using this format:
        {mermaid_template}

        If there is no process flow or workflow, respond 'No process flow detected.'
        """
        
        try:
            response = self._invoke_completion(
                messages=[
                    {"role": "system", "content": "You are an expert business process analyst who creates clear process maps from meeting transcripts."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.3
            )
            # usage = response.usage
            # if usage:
            #     log_openai_usage(
            #         module_name=__name__,
            #         model=get_chat_model_name(),
            #         prompt_tokens=usage.prompt_tokens,
            #         completion_tokens=usage.completion_tokens,
            #         function_name="_generate_process_map",
            #         session_id=self.session_guid
            #     )

            if response:
                token_tracker.track_response(response=response, auth_token=self.auth_token, model=self._litellm_model)

            return response.choices[0].message.content
        except Exception as e:
            logging.info(f"Error generating process map: {e}")
            return "## Process Map\n\nError generating process map. Please check OpenAI configuration."
        
    def _generate_narrative_documentation(self) -> Dict[str, Any]:
        """
        Generate a complete narrative-based documentation from the speech transcript with
        screenshots integrated at appropriate points.
        
        Returns:
            Dictionary containing structured documentation content
        """
        logging.info(f"DEBUG: Generating narrative documentation with {len(self.screenshots)} screenshots")
        if not self.use_ai:
            # Fall back to basic structure if AI is not available
            return {
                "title": self.title,
                "introduction": self.description,
                "sections": [
                    {
                        "title": "Product Overview",
                        "content": "This document provides an overview of the product demonstrated in the video.",
                        "subsections": []
                    }
                ]
            }
        
        # Use the speech segments provided during initialization
        speech_segments = self.speech_segments.copy()
        
        # If no speech segments were provided, try to extract them from session state
        if not speech_segments:
            try:
                # Access all speech segments from the app's session state
                import streamlit as st
                if 'speech_timestamps' in st.session_state and st.session_state.speech_timestamps:
                    logging.info(f"Found {len(st.session_state.speech_timestamps)} speech segments in session state")
                    speech_segments = st.session_state.speech_timestamps
            except Exception as e:
                logging.info(f"Error accessing session state: {e}")
        
        # If still no speech segments, try to extract from screenshot reasons
        if not speech_segments:
            logging.info("No speech segments provided, extracting from screenshot reasons")
            # Extract from screenshot reasons
            for _, ts, reason in self.screenshots:
                if "Keyword trigger" in reason and ":" in reason:
                    # Extract the speech text from the reason
                    parts = reason.split(":", 1)
                    if len(parts) > 1:
                        speech_text = parts[1].strip()
                        speech_segments.append((ts, speech_text))
                elif "AI detected" in reason:
                    # Get any additional speech context from AI detection
                    parts = reason.split(":", 1)
                    if len(parts) > 1:
                        speech_segments.append((ts, parts[1].strip()))
        
        # If we still don't have enough segments, create a fallback transcript from the video title
        if not speech_segments or len(speech_segments) < 2:
            logging.info("Not enough speech segments, creating a basic document structure")
            # Create a structured document with limited speech data
            return {
                "title": self.title,
                "introduction": f"Documentation for {self.title}",
                "sections": [
                    {
                        "title": "Product Overview",
                        "content": f"This document provides a visual overview of {self.title}. " +
                                   f"The screenshots capture key moments and features from the video demonstration.",
                        "screenshot_timestamps": [screenshot[1] for screenshot in self.screenshots[:3]],
                        "subsections": [
                            {
                                "title": "Key Features",
                                "content": "The product demonstrates various features and user interactions.",
                                "screenshot_timestamps": [screenshot[1] for screenshot in self.screenshots[3:6] if len(self.screenshots) > 3]
                            },
                            {
                                "title": "User Interface",
                                "content": "The interface provides an intuitive way to navigate and access functionality.",
                                "screenshot_timestamps": [screenshot[1] for screenshot in self.screenshots[6:] if len(self.screenshots) > 6]
                            }
                        ]
                    }
                ]
            }
            
        # Sort speech segments by timestamp
        speech_segments.sort(key=lambda x: x[0])
        
        # Prepare the full transcript for OpenAI analysis
        full_transcript = ""
        for timestamp, text in speech_segments:
            full_transcript += f"[{timestamp:.2f}s] {text}\n"
        
        # Enhance screenshot reasons with AI analysis if available
        enhanced_screenshots = self._enhance_screenshot_reasons(speech_segments)
        
        # Create enhanced screenshot context with 5-second chunking and 30-second speech context
        screenshot_context = "\n🎯 SCREENSHOTS WITH DETAILED CONTEXT:\n"
        
        for enhanced_img, ts, enhanced_reason in enhanced_screenshots:
            # Get 30-second speech context around this screenshot (15 seconds before and after)
            context_start = max(0, ts - 15)
            context_end = ts + 15
            
            speech_context = []
            for speech_ts, speech_text in speech_segments:
                if context_start <= speech_ts <= context_end:
                    time_diff = speech_ts - ts
                    if time_diff < 0:
                        speech_context.append(f"  [{time_diff:.1f}s before] {speech_text}")
                    elif time_diff > 0:
                        speech_context.append(f"  [+{time_diff:.1f}s after] {speech_text}")
                    else:
                        speech_context.append(f"  [EXACT TIME] {speech_text}")
            
            screenshot_context += f"\n📷 Screenshot at {ts:.2f}s: {enhanced_reason}\n"
            if speech_context:
                screenshot_context += "   Speech Context (30-second window):\n"
                for context_line in speech_context[:6]:  # Limit to 6 most relevant lines
                    screenshot_context += f"{context_line}\n"
            else:
                screenshot_context += "   No speech detected in 30-second window\n"
        
        screenshot_context += "\n📍 CONTEXTUAL PLACEMENT INSTRUCTIONS:\n"
        screenshot_context += "- Use the speech context to understand what was being discussed at each screenshot moment\n"
        screenshot_context += "- Place screenshots in sections where the speech content aligns with the section topic\n"
        screenshot_context += "- Each screenshot should appear exactly ONCE in the most contextually relevant section\n"
        screenshot_context += "- Consider both the screenshot reason AND the surrounding speech when choosing placement\n"
        screenshot_context += "- Create sections that naturally flow with the speech narrative\n"
        screenshot_context += "- Skip screenshots that don't fit naturally into any section"
        
        # Prepare analysis prompt based on document type
        if self.document_type == "user_story_generator":
            # Meeting Summary with metadata
            metadata_info = ""
            if self.meeting_participants:
                metadata_info += f"\nMeeting Attendees: {', '.join(self.meeting_participants)}"
            if self.meeting_highlights:
                metadata_info += f"\nKey Discussion Points: {', '.join(self.meeting_highlights)}"
            if self.meeting_duration_minutes:
                hours = int(self.meeting_duration_minutes // 60)
                minutes = int(self.meeting_duration_minutes % 60)
                duration = f"{hours}h {minutes}m" if hours > 0 else f"{minutes}m"
                metadata_info += f"\nMeeting Duration: {duration}"
            
            prompt = f"""
            I need you to analyze the transcript from a requirements discussion or a meeting and generate well-structured user stories.

            TRANSCRIPT:
            {full_transcript}
            {screenshot_context}

            This is a USER STORY GENERATOR document that should:
            1. Identify and extract user requirements from the discussion
            2. Format each requirement as a proper user story following the format "As a [type of user], I want [goal] so that [benefit]"
            3. Add acceptance criteria for each user story
            4. Group related user stories into epics or features
            5. Identify timestamps where requirements are being discussed (for screenshot placement)
            6. Assign priority levels to user stories when possible (High/Medium/Low)
            8. There might be other relevant information in the discussion as well. Capture all useful information from the discussion.
            9. Donot generate additional own content. Everything has to be based on the transcript.


            Format your response as a JSON object with the following structure:
            {{
                "title": "User Stories: [Project Name]",
                "introduction": "Collection of user stories derived from requirements discussions for [Project Name]",
                "sections": [
                    {{
                        "title": "Epic/Feature Name",
                        "content": "Overview of this group of related user stories",
                        "screenshot_timestamps": [list of timestamps where requirements for this epic were discussed],
                        "subsections": [
                            {{
                                "title": "User Story: [Brief Story Title]",
                                "content": "As a [user type], I want [goal] so that [benefit]\\n\\nAcceptance Criteria:\\n- Criterion 1\\n- Criterion 2\\n\\nPriority: [priority level]",
                                "screenshot_timestamps": [list of timestamps relevant to this specific user story]
                            }}
                        ]
                    }}
                ]
            }}
            """
        try:
            from openai.types.chat import ChatCompletionSystemMessageParam, ChatCompletionUserMessageParam
            
            # Prepare system message based on document type
            
            if self.document_type == "user_story_generator":
                system_message: ChatCompletionSystemMessageParam = {
                    "role": "system", 
                    "content": "You are a requirements analysis expert. Your task is to extract user stories from discussions and format them in a structured way with clear acceptance criteria."
                }
            
            user_message: ChatCompletionUserMessageParam = {
                "role": "user", 
                "content": prompt
            }
            
            # Call OpenAI API
            try:
                if not self.use_ai or self._litellm_model is None:
                    raise ValueError("Model name not configured properly")

                response = self._invoke_completion(
                    messages=[system_message, user_message],
                    temperature=0.7,
                    response_format={"type": "json_object"}
                )

                # usage = response.usage
                # if usage:
                #     log_openai_usage(
                #         module_name=__name__,
                #         model=self.model or "",
                #         prompt_tokens=usage.prompt_tokens,
                #         completion_tokens=usage.completion_tokens,
                #         function_name="_generate_narrative_documentation",
                #         session_id=self.session_guid
                #     )
                if response:
                    token_tracker.track_response(response=response, auth_token=self.auth_token, model=self._litellm_model)
                logging.info(f"API call successful for narrative document generation")
            except Exception as e:
                logging.info(f"API call error details: {str(e)}")
                raise
            
            # Parse generated documentation structure
            content = response.choices[0].message.content
            if content:
                cleaned_content = self._clean_json_response(content)
                try:
                    doc_structure = json.loads(cleaned_content)
                    
                    # Debug the document structure
                    logging.info(f"DEBUG: Document structure received from AI")
                    logging.info(f"DEBUG: Document title: {doc_structure.get('title', 'No title')}")
                    logging.info(f"DEBUG: Document has {len(doc_structure.get('sections', []))} sections")
                    
                    # Check if the sections have screenshot_timestamps
                    for i, section in enumerate(doc_structure.get('sections', [])):
                        has_timestamps = 'screenshot_timestamps' in section
                        timestamp_count = len(section.get('screenshot_timestamps', []))
                        logging.info(f"DEBUG: Section {i+1} has timestamps: {has_timestamps}, count: {timestamp_count}")
                    
                    # Clean screenshot verification - only use AI assignments, no duplicates
                    all_screenshot_timestamps = [ts for _, ts, _ in self.screenshots]
                    logging.info(f"DEBUG: Have {len(all_screenshot_timestamps)} total screenshots available")
                    
                    # Count total AI assignments and validate they exist
                    total_ai_assignments = 0
                    invalid_count = 0
                    available_timestamps = set(s[1] for s in self.screenshots)
                    
                    def find_closest_timestamp(target_ts, available_ts_list, tolerance=2.0):
                        """Find the closest available timestamp within tolerance"""
                        try:
                            # Clean the target timestamp - remove 's' suffix if present
                            if isinstance(target_ts, str):
                                target_ts = target_ts.rstrip('s')
                            target_ts = float(target_ts)  # Ensure target is a number
                            
                            for avail_ts in available_ts_list:
                                avail_ts = float(avail_ts)  # Ensure available timestamp is a number
                                if abs(avail_ts - target_ts) <= tolerance:
                                    return avail_ts
                        except (ValueError, TypeError) as e:
                            logging.info(f"DEBUG: Timestamp conversion error: {e}, target: {target_ts}")
                        return None
                    
                    available_timestamps_list = list(available_timestamps)
                    
                    for section in doc_structure.get('sections', []):
                        total_ai_assignments += len(section.get('screenshot_timestamps', []))
                        # Match timestamps with tolerance for precision differences
                        if 'screenshot_timestamps' in section:
                            valid_timestamps = []
                            for ts in section['screenshot_timestamps']:
                                matched_ts = find_closest_timestamp(ts, available_timestamps_list)
                                if matched_ts is not None:
                                    valid_timestamps.append(matched_ts)
                                    logging.info(f"DEBUG: Matched AI timestamp {ts}s to available {matched_ts}s")
                                else:
                                    invalid_count += 1
                                    logging.info(f"DEBUG: No match for AI timestamp {ts}s")
                            section['screenshot_timestamps'] = valid_timestamps
                        
                        for subsection in section.get('subsections', []):
                            total_ai_assignments += len(subsection.get('screenshot_timestamps', []))
                            # Match timestamps with tolerance for precision differences
                            if 'screenshot_timestamps' in subsection:
                                valid_timestamps = []
                                for ts in subsection['screenshot_timestamps']:
                                    matched_ts = find_closest_timestamp(ts, available_timestamps_list)
                                    if matched_ts is not None:
                                        valid_timestamps.append(matched_ts)
                                        logging.info(f"DEBUG: Matched AI timestamp {ts}s to available {matched_ts}s")
                                    else:
                                        invalid_count += 1
                                        logging.info(f"DEBUG: No match for AI timestamp {ts}s")
                                subsection['screenshot_timestamps'] = valid_timestamps
                    
                    # Count final assignments
                    final_assignments = 0
                    for section in doc_structure.get('sections', []):
                        final_assignments += len(section.get('screenshot_timestamps', []))
                        for subsection in section.get('subsections', []):
                            final_assignments += len(subsection.get('screenshot_timestamps', []))
                    
                    logging.info(f"DEBUG: AI made {total_ai_assignments} total assignments for {len(all_screenshot_timestamps)} screenshots")
                    logging.info(f"DEBUG: Removed {invalid_count} invalid timestamps")
                    logging.info(f"DEBUG: Final valid assignments: {final_assignments}")
                    logging.info(f"DEBUG: Allowing multiple placements for better context")
                    
                    # Add missing questions section if enabled
                    if self.generate_missing_questions and self.use_ai:
                        logging.info("DEBUG: Generating missing questions section")
                        missing_questions = self._generate_missing_questions()
                        if missing_questions and missing_questions.strip():
                            doc_structure.setdefault('sections', []).append({
                                "title": "Missing Questions for Next Meeting",
                                "content": missing_questions,
                                "screenshot_timestamps": [],
                                "subsections": []
                            })
                    
                    # Add process map section if enabled
                    if self.generate_process_map and self.use_ai:
                        logging.info("DEBUG: Generating process map section")
                        process_map = self._generate_process_map()
                        if process_map and process_map.strip() and "No process workflow identified" not in process_map:
                            # Extract mermaid code and create visual diagram
                            diagram_image = None
                            mermaid_code = None
                            if "```mermaid" in process_map:
                                mermaid_start = process_map.find("```mermaid")
                                mermaid_end = process_map.find("```", mermaid_start + 10) + 3  # End of the closing ```
                                if mermaid_end > mermaid_start:
                                    mermaid_code = process_map[mermaid_start + 10:mermaid_end - 3].strip()  # Extract Mermaid code
                                    logging.info(f"DEBUG: Extracted mermaid code: {mermaid_code[:100]}...")
                                    generator = MermaidDiagramGenerator()
                                    diagram_image = generator.create_mermaid_diagram(
                                                            mermaid_code=mermaid_code,
                                                            deployment_name="gpt-4o",  # Your Azure OpenAI deployment name
                                                            enable_ai_fix=True,
                                                            max_ai_attempts = 3
                                                            )
                                    if diagram_image:
                                        if mermaid_end > mermaid_start:
                                         # Remove the entire Mermaid code block, including delimiters
                                            process_map =  process_map[:mermaid_start] + process_map[mermaid_end:]
                                        logging.info("DEBUG: Successfully created mermaid diagram image")
                                    else:
                                        logging.info("DEBUG: Failed to create mermaid diagram image")
                            
                            doc_structure.setdefault('sections', []).append({
                                "title": "Process Map",
                                "content": process_map,
                                "diagram_image": diagram_image,
                                "mermaid_code" : mermaid_code,
                                "screenshot_timestamps": [],
                                "subsections": []
                            })
                    
                    return doc_structure
                    
                except json.JSONDecodeError as e:
                    logging.info(f"Error parsing JSON response: {e}")
                    logging.info(f"Raw content: {content}")
            
            # Return default structure if parsing failed
            default_structure = {
                "title": self.title,
                "introduction": "This document provides an overview of the product demonstrated in the video.",
                "sections": [
                    {
                        "title": "Product Features",
                        "content": "The product demo showcases several key features and capabilities.",
                        "screenshot_timestamps": [s[1] for s in self.screenshots[:3]],
                        "subsections": []
                    },
                    {
                        "title": "Usage Examples",
                        "content": "The following screenshots illustrate how to use key features.",
                        "screenshot_timestamps": [s[1] for s in self.screenshots[3:] if len(self.screenshots) > 3],
                        "subsections": []
                    }
                ]
            }
            
            logging.info("DEBUG: Using default document structure with screenshot timestamps")
            return default_structure
            
        except Exception as e:
            logging.info(f"Error generating narrative documentation: {e}")
            return {
                "title": self.title,
                "introduction": self.description,
                "sections": [
                    {
                        "title": "Video Documentation",
                        "content": "Please refer to the screenshots below for an overview of the product.",
                        "screenshot_timestamps": [s[1] for s in self.screenshots[:4]],
                        "subsections": []
                    },
                    {
                        "title": "Additional Screenshots",
                        "content": "These screenshots provide further details about the content shown in the video.",
                        "screenshot_timestamps": [s[1] for s in self.screenshots[4:] if len(self.screenshots) > 4],
                        "subsections": []
                    }
                ]
            }
    def _generate_section_descriptions(self) -> Dict[float, str]:
        """
        Generate detailed descriptions for each screenshot section using OpenAI
        
        Returns:
            Dictionary mapping timestamp to detailed description
        """
        # This is now a compatibility method for backwards compatibility
        # We'll use a simpler approach now that we're focused on narrative documentation
        return {ts: self._format_reason(reason) for _, ts, reason in self.screenshots}
    
    def _format_reason(self, reason: str) -> str:
        """Format detection reason into a readable description"""
        
        if "Scene change" in reason:
            return f"New screen or view detected. {reason}"
        elif "AI Detected" in reason:
            # Extract the actual reason from AI detection
            parts = reason.split(":", 1)
            if len(parts) > 1:
                return parts[1].strip()
            return reason
        elif "Keyword trigger" in reason:
            return f"Important interaction detected: {reason.split(':', 1)[1].strip() if ':' in reason else reason}"
        elif "Mouse Click" in reason:
            return "User interaction detected - mouse click or selection action"
        elif "Text Change" in reason:
            return "Content update detected - significant text or data change"
        else:
            return reason
    
    def _prepare_screenshot_tracking(self) -> Dict[str, Any]:
        """
        Prepare screenshot tracking system to prevent duplicate screenshots 
        in document generation.
        
        Returns:
            Dictionary containing screenshot tracking data:
            - screenshot_map: Map of timestamps to (img, reason) pairs
            - used_screenshots: Set to track which screenshots have been used
            - find_screenshot: Function to find the best unused screenshot for a timestamp
        """
        # Create a map of timestamps to screenshots
        screenshot_map = {}
        logging.info(f"DEBUG: Total screenshots to process: {len(self.screenshots)}")
        
        # Check if screenshots are valid before adding to map
        for i, (img, timestamp, reason) in enumerate(self.screenshots):
            if img is None:
                logging.info(f"DEBUG: Screenshot {i} has None image at timestamp {timestamp}")
                continue
                
            try:
                # Verify the image is valid
                img_format = img.format
                img_size = img.size
                logging.info(f"DEBUG: Screenshot {i} at {timestamp}s is valid: format={img_format}, size={img_size}")
                screenshot_map[timestamp] = (img, reason)
            except Exception as e:
                logging.info(f"DEBUG: Error with screenshot {i} at {timestamp}s: {str(e)}")
        
        logging.info(f"DEBUG: Valid screenshots in map: {len(screenshot_map)}")
        
        # Set to track which timestamps have been used already in the document
        used_timestamps = set()
        
        # Function to find an unused screenshot closest to the requested timestamp
        def find_closest_unused_screenshot(target_time, max_time_diff=8.0):
            """Find the best unused screenshot near target_time"""
            logging.info(f"DEBUG: Looking for screenshot near {target_time}s")
            
            if not screenshot_map:
                logging.info("DEBUG: No screenshots in map")
                return None, None
                
            # Handle string timestamps
            if isinstance(target_time, str):
                try:
                    target_time = float(target_time)
                    logging.info(f"DEBUG: Converted string timestamp to float: {target_time}")
                except ValueError:
                    logging.info(f"DEBUG: Could not convert timestamp '{target_time}' to float")
                    return None, None
            
            # Check for exact or very close match (handle floating point precision)
            for ts in screenshot_map.keys():
                if ts not in used_timestamps:
                    # Check for exact match or within 0.1 second tolerance for floating point precision
                    if abs(ts - target_time) <= 0.1:
                        used_timestamps.add(ts)
                        logging.info(f"DEBUG: Found exact match at {ts}s")
                        return screenshot_map[ts]
            
            # Find closest unused screenshot within time window
            closest_ts = None
            min_diff = float('inf')
            
            for ts in screenshot_map.keys():
                if ts not in used_timestamps:
                    diff = abs(ts - target_time)
                    if diff < min_diff and diff <= max_time_diff:
                        min_diff = diff
                        closest_ts = ts
            
            if closest_ts is not None:
                used_timestamps.add(closest_ts)
                logging.info(f"DEBUG: Found closest match at {closest_ts}s (diff: {min_diff:.2f}s)")
                return screenshot_map[closest_ts]
            
            logging.info(f"DEBUG: No suitable screenshot found near {target_time}s")
            return None, None
        
        return {
            "screenshot_map": screenshot_map,
            "used_timestamps": used_timestamps,
            "find_screenshot": find_closest_unused_screenshot
        }
    
    def generate_docx(self, output_path: str = "") -> str:
        """
        Generate a Word document with the screenshots and descriptions
        
        Args:
            output_path: Path to save the document (optional)
            
        Returns:
            Path to the generated document
        """
        # Create a new document
        doc = Document()
        
        # Set document title and add heading
        doc.add_heading(self.title, 0)
        
        # Add document description
        doc.add_paragraph(self.description)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Generate narrative documentation if AI is available
        if self.use_ai and OPENAI_AVAILABLE:
            doc_structure = self._generate_narrative_documentation()
            
            # Check if we have a saved process diagram file
            import os
            temp_diagram_path = "temp_process_diagram.png"
            
            # Use your regex pattern approach to find and embed mermaid diagrams
            import re
            pattern = r'```mermaid\s+(.*?)\s+```'
            
            # Process each section from the narrative structure
            for section in doc_structure.get('sections', []):
                # Add section heading
                doc.add_heading(section['title'], level=1)
                
                # Add section content with proper formatting
                content = section.get('content', '')
                if content:
                    # First, extract any mermaid diagrams using your pattern
                    parts = re.split(pattern, content, flags=re.DOTALL)
                    
                    for i, part in enumerate(parts):
                        part = part.strip()
                        if not part:
                            continue
                            
                        # Check if this part looks like mermaid code
                        if 'flowchart' in part or '-->' in part:
                            # This is mermaid code, create and add diagram
                            try:
                                mermaid_image = self._create_mermaid_diagram(part)
                                if mermaid_image:
                                    # Save temporarily using your approach
                                    temp_path = f"temp_mermaid_{i}.png"
                                    mermaid_image.save(temp_path, format='PNG')
                                    optimal_width, optimal_height = self._calculate_optimal_image_size(mermaid_image)

                                    # Add to document
                                    doc.add_paragraph("Process Flow Diagram:", style='Heading 3')
                                    doc.add_picture(temp_path, width=Inches(optimal_width))
                                    logging.info(f"DEBUG: Added mermaid diagram from content: {temp_path}")
                                    
                                    # Clean up temp file
                                    if os.path.exists(temp_path):
                                        os.remove(temp_path)
                                        
                            except Exception as e:
                                logging.info(f"DEBUG: Error processing mermaid from content: {e}")
                                doc.add_paragraph("[Process diagram could not be loaded]")
                        else:
                            # This is regular text content
                            if '**' in part:
                                lines = part.split('\n')
                                for line in lines:
                                    line = line.strip()
                                    if not line:
                                        continue
                                    # Handle bold headings (**TEXT**)
                                    if line.startswith('**') and line.endswith('**'):
                                        heading_text = line.strip('*')
                                        para = doc.add_paragraph()
                                        para.add_run(heading_text).bold = True
                                    # Handle bullet points
                                    elif line.startswith('• ') or line.startswith('- '):
                                        bullet_text = line[2:]
                                        doc.add_paragraph(bullet_text, style='List Bullet')
                                    # Regular paragraphs
                                    else:
                                        doc.add_paragraph(line)
                            else:
                                paragraphs = part.split('\n\n')
                                for para in paragraphs:
                                    if para.strip():
                                        doc.add_paragraph(para.strip())
                
                # Add screenshots for this section if any
                screenshot_tracking = self._prepare_screenshot_tracking()
                for timestamp in section.get('screenshot_timestamps', []):
                    screenshot_result = screenshot_tracking['find_screenshot'](timestamp)
                    if screenshot_result:
                        img, actual_timestamp, reason = screenshot_result
                        
                        # Add screenshot with caption
                        from ..utils.media_utils import format_timestamp
                        caption = f"Screenshot at {format_timestamp(actual_timestamp)}"
                        caption_para = doc.add_paragraph()
                        caption_para.add_run(caption).italic = True
                        
                        # Add the image
                        img_buffer = BytesIO()
                        img.save(img_buffer, format="PNG")
                        img_buffer.seek(0)
                        
                        try:
                            doc.add_picture(img_buffer, width=Inches(6))
                        except Exception as e:
                            logging.info(f"Error adding screenshot to DOCX: {e}")
                            doc.add_paragraph(f"[Image could not be loaded: {e}]")
                
                # Add process diagram if this is a process map section
                if section['title'] == 'Process Map' and section.get('diagram_image'):
                    try:
                        # Save diagram to a temporary file
                        temp_path = "temp_process_diagram.png"
                        section['diagram_image'].save(temp_path, format='PNG')
                        optimal_width, optimal_height = self._calculate_optimal_image_size(section['diagram_image'])

                        # Add to document
                        doc.add_paragraph("Process Flow Diagram:", style='Heading 3')
                        doc.add_picture(temp_path, width=Inches(optimal_width))
                        logging.info(f"DEBUG: Added process diagram to DOCX: {temp_path}")
                        
                        # Clean up temp file
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    except Exception as e:
                        logging.info(f"Error adding process diagram to DOCX: {e}")
                        doc.add_paragraph("[Process diagram could not be loaded]")
                
                # Add page break between sections
                doc.add_page_break()
        
        else:
            # Fallback to basic format if AI is not available
            doc.add_heading('Contents', level=1)
            doc.add_paragraph('Screenshots and Descriptions')
            doc.add_page_break()
            
            # Generate descriptions for each screenshot
            descriptions = self._generate_section_descriptions()
            
            # Add each screenshot with its description
            for i, (img, timestamp, reason) in enumerate(self.screenshots):
                # Add section heading with timestamp
                from ..utils.media_utils import format_timestamp
                heading_text = f"Screenshot {i+1}: {format_timestamp(timestamp)}"
                doc.add_heading(heading_text, level=2)
                
                # Add description
                description = descriptions.get(timestamp, self._format_reason(reason))
                if isinstance(description, dict):
                    description_text = str(description)
                else:
                    description_text = description
                doc.add_paragraph(description_text)
                
                # Add the image
                img_buffer = BytesIO()
                img.save(img_buffer, format="PNG")
                img_buffer.seek(0)
                
                try:
                    doc.add_picture(img_buffer, width=Inches(6))
                except Exception as img_err:
                    logging.info(f"Error adding image to DOCX: {img_err}")
                    doc.add_paragraph(f"[Image placeholder - could not load image: {img_err}]")
                
                # Add detection method info
                detection_para = doc.add_paragraph("Detection method: ")
                if "Scene change" in reason:
                    detection_para.add_run("Scene Change Analysis").bold = True
                elif "AI Detected" in reason:
                    detection_para.add_run("AI Speech Analysis").bold = True
                elif "Keyword trigger" in reason:
                    detection_para.add_run("Speech Keyword Detection").bold = True
                elif "Mouse Click" in reason:
                    detection_para.add_run("User Interaction Detection").bold = True
                elif "Text Change" in reason:
                    detection_para.add_run("Content Change Detection").bold = True
                else:
                    detection_para.add_run("Other").bold = True
                
                # Add a page break if not the last screenshot
                if i < len(self.screenshots) - 1:
                    doc.add_page_break()
        
        # Determine output path if not provided
        if not output_path:
            base_name = os.path.basename(self.video_path).split('.')[0]
            output_path = f"{base_name}_documentation.docx"
        
        # Save the document
        doc.save(output_path)
        
        return output_path
    def generate_pdf(self, output_path: str = "") -> str:
        """
        Generate a PDF document with the screenshots and descriptions
        
        Args:
            output_path: Path to save the document (optional)
            
        Returns:
            Path to the generated document
        """
        import os
        
        # Determine output path if not provided
        if not output_path:
            base_name = os.path.basename(self.video_path).split('.')[0]
            output_path = f"{base_name}_documentation.pdf"
        
        # Create PDF document
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        heading1_style = styles['Heading1']
        heading2_style = styles['Heading2']
        normal_style = styles['Normal']
        bold_style = ParagraphStyle('Bold', parent=normal_style, fontName='Helvetica-Bold')
        
        # Create content elements
        elements = []
        
        # Add title
        elements.append(Paragraph(self.title, title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Add description
        elements.append(Paragraph(self.description, normal_style))
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Generate narrative documentation structure
        if self.use_ai and OPENAI_AVAILABLE:
            doc_structure = self._generate_narrative_documentation()
            
            # Add sections from the narrative structure
            for section in doc_structure.get('sections', []):
                # Add section heading
                elements.append(Paragraph(section['title'], heading1_style))
                elements.append(Spacer(1, 0.2*inch))
                
                # Add section content with proper formatting
                content = section.get('content', '')
                if content:
                    if '**' in content:
                        formatted_content = self._format_content_for_pdf(content, normal_style, bold_style)
                        elements.extend(formatted_content)
                    else:
                        elements.append(Paragraph(content, normal_style))
                    elements.append(Spacer(1, 0.2*inch))
                
                # Add screenshots for this section if any
                screenshot_tracking = self._prepare_screenshot_tracking()
                for timestamp in section.get('screenshot_timestamps', []):
                    screenshot_result = screenshot_tracking['find_screenshot'](timestamp)
                    if screenshot_result and screenshot_result[0]:
                        img, actual_timestamp, reason = screenshot_result
                        
                        # Add screenshot with caption
                        from ..utils.media_utils import format_timestamp
                        caption = f"Screenshot at {format_timestamp(actual_timestamp)}"
                        elements.append(Paragraph(caption, 
                                                ParagraphStyle('caption', parent=normal_style, 
                                                            fontSize=10, textColor=blue)))
                        
                        # Add the image
                        img_buffer = BytesIO()
                        img.save(img_buffer, format="PNG")
                        img_buffer.seek(0)
                        
                        try:
                            img_for_pdf = Image(img_buffer, width=5*inch, height=3*inch)
                            elements.append(img_for_pdf)
                        except Exception as e:
                            logging.info(f"Error adding screenshot to PDF: {e}")
                            elements.append(Paragraph(f"[Image could not be loaded: {e}]", normal_style))
                        
                        elements.append(Spacer(1, 0.3*inch))
                
                # Add visual diagram if this is a process map section
                if section['title'] == 'Process Map' and section.get('diagram_image'):
                    try:
                        # Verify input image
                        if not isinstance(section['diagram_image'], PILImage.Image):
                            raise ValueError("diagram_image is not a valid PIL Image object")

                        # Convert PIL image to bytes for PDF
                        diagram_buffer = BytesIO()
                        section['diagram_image'].convert('RGB').save(diagram_buffer, format='PNG', dpi=(100, 100))
                        diagram_buffer.seek(0)
                        logging.info("image_size: ", section['diagram'].size)
                        # Calculate optimal dimensions
                        max_width_inches = 7.0  # Adjust based on page layout
                        max_height_inches = 5.0  # Adjust based on page layout
                        width_inches, height_inches = self._calculate_optimal_image_size_pdf(
                            section['diagram_image'], max_width_inches, max_height_inches
                        )

                        # Add elements to the PDF
                        elements.append(Paragraph(
                            "Process Flow Diagram:",
                            ParagraphStyle(
                                'diagram_header',
                                parent=normal_style,
                                fontSize=12,
                                fontName='Helvetica-Bold'
                            )
                        ))
                        elements.append(Spacer(1, 0.1*inch))

                        # Create Image with calculated dimensions
                        diagram_img = Image(
                            diagram_buffer,
                            width=width_inches*inch,
                            height=height_inches*inch,
                            kind='proportional'  # Ensure aspect ratio is maintained
                        )
                        # Disable any automatic resizing
                        diagram_img.hAlign = 'CENTER'
                        diagram_img._restrictSize(width_inches*inch, height_inches*inch)

                        elements.append(diagram_img)
                        elements.append(Spacer(1, 0.3*inch))
                    except Exception as e:
                        logging.info(f"Error adding process diagram to PDF: {e}")
                        elements.append(Paragraph(
                            f"[Process diagram could not be loaded: {e}]",
                            normal_style
                        ))
                elements.append(Spacer(1, 0.4*inch))
        else:
            # Fallback to original screenshot-based format
            elements.append(Paragraph("Screenshots and Descriptions", heading1_style))
            elements.append(Spacer(1, 0.5*inch))
            
            # Generate descriptions for each screenshot
            descriptions = self._generate_section_descriptions()
            
            # Add each screenshot with its description
            for i, (img, timestamp, reason) in enumerate(self.screenshots):
                # Add section heading with timestamp
                from ..utils.media_utils import format_timestamp
                heading_text = f"Screenshot {i+1}: {format_timestamp(timestamp)}"
                elements.append(Paragraph(heading_text, heading2_style))
                elements.append(Spacer(1, 0.1*inch))
                
                # Add description
                description = descriptions.get(timestamp, self._format_reason(reason))
                if isinstance(description, dict):
                    description_text = str(description)
                else:
                    description_text = description
                elements.append(Paragraph(description_text, normal_style))
                elements.append(Spacer(1, 0.1*inch))
                
                # Add the image
                img_buffer = BytesIO()
                img.save(img_buffer, format="PNG")
                img_buffer.seek(0)
                
                try:
                    img_for_pdf = Image(img_buffer, width=6*inch, height=3.5*inch)
                    elements.append(img_for_pdf)
                except Exception as img_err:
                    logging.info(f"Error adding image to PDF: {img_err}")
                    elements.append(Paragraph(f"[Image placeholder - could not load image: {img_err}]", normal_style))
                elements.append(Spacer(1, 0.1*inch))
                
                # Add detection method info
                if "Scene change" in reason:
                    detection_text = "Detection method: Scene Change Analysis"
                elif "AI Detected" in reason:
                    detection_text = "Detection method: AI Speech Analysis"
                elif "Keyword trigger" in reason:
                    detection_text = "Detection method: Speech Keyword Detection"
                elif "Mouse Click" in reason:
                    detection_text = "Detection method: User Interaction Detection"
                elif "Text Change" in reason:
                    detection_text = "Detection method: Content Change Detection"
                else:
                    detection_text = "Detection method: Other"
                    
                elements.append(Paragraph(detection_text, 
                                        ParagraphStyle('detection', 
                                                    parent=normal_style, 
                                                    fontName='Helvetica-Bold')))
                elements.append(Spacer(1, 0.5*inch))
        
        # Build the document
        doc.build(elements)
        
        # Save the PDF to the output path
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())
        
        # Clean up temporary image files
        self._cleanup_temp_files()
        
        return output_path
    def _cleanup_temp_files(self):
        """Clean up any temporary files created during document generation"""
        # Temporary files are cleaned up immediately after use
        # This method is kept for future extension if needed
        pass
    
    def get_document_bytes(self, doc_type="pdf") -> bytes:
        """
        Generate a document in memory and return its bytes
        
        Args:
            doc_type: Type of document to generate ("pdf" or "docx")
            
        Returns:
            Bytes of the generated document
        """
        import io
        
        logging.info(f"DEBUG: Starting document generation with {len(self.screenshots)} screenshots")
        valid_screenshots = []
        doc_structure = self._generate_narrative_documentation()

        if self.include_screenshots:
            for i, (img, timestamp, reason) in enumerate(self.screenshots):
                if img is not None:
                    valid_screenshots.append((img, timestamp, reason))
                else:
                    logging.info(f"WARNING: Screenshot {i} has None image at timestamp {timestamp}")
            
            if not valid_screenshots:
                logging.info("ERROR: No valid screenshots available for document generation")
                if doc_type.lower() == "pdf":
                    buffer = BytesIO()
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    styles = getSampleStyleSheet()
                    elements = []
                    elements.append(Paragraph(self.title, styles['Title']))
                    elements.append(Paragraph("No valid screenshots were found for document generation.", styles['Normal']))
                    doc.build(elements)
                    buffer.seek(0)
                    return buffer.getvalue()
                else:
                    from docx import Document
                    doc = Document()
                    doc.add_heading(self.title, 0)
                    doc.add_paragraph("No valid screenshots were found for document generation.")
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    return doc_buffer.getvalue()
            
            self.screenshots = valid_screenshots
            logging.info(f"DEBUG: Proceeding with {len(self.screenshots)} valid screenshots")
            
        if doc_type.lower() == "pdf":
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            heading1_style = styles['Heading1']
            heading2_style = styles['Heading2']
            normal_style = styles['Normal']
            bold_style = ParagraphStyle('Bold', parent=normal_style, fontName='Helvetica-Bold')
            
            # Create content elements
            elements = []
            
            # Add title
            elements.append(Paragraph(self.title, title_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add description
            elements.append(Paragraph(self.description, normal_style))
            elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add table of contents heading
            elements.append(Paragraph("Contents", heading1_style))
            if self.include_screenshots:
                elements.append(Paragraph("Screenshots and Descriptions", normal_style))
                elements.append(Spacer(1, 0.5*inch))
            
            # Generate descriptions for each screenshot
            descriptions = self._generate_section_descriptions()
            
            # IMPORTANT: Make sure the document structure has screenshot timestamps
            # Add screenshots to sections if they don't have any
            if doc_structure and 'sections' in doc_structure:
                # Count sections with screenshot timestamps
                sections_with_timestamps = sum(1 for s in doc_structure['sections'] 
                                              if 'screenshot_timestamps' in s and s['screenshot_timestamps'])
                if self.include_screenshots:
                    if sections_with_timestamps == 0:
                        logging.info("DEBUG: No sections have screenshot timestamps, distributing screenshots")
                        all_timestamps = [ts for _, ts, _ in self.screenshots]
                        num_sections = len(doc_structure['sections'])
                        screenshots_per_section = max(1, len(all_timestamps) // num_sections)
                        
                        for i, section in enumerate(doc_structure['sections']):
                            start_idx = i * screenshots_per_section
                            end_idx = min(start_idx + screenshots_per_section, len(all_timestamps))
                            section['screenshot_timestamps'] = all_timestamps[start_idx:end_idx]
                            logging.info(f"DEBUG: Added {len(section['screenshot_timestamps'])} screenshots to section {i+1}")
                
            # Add document title
            elements.append(Paragraph(doc_structure.get("title", self.title), heading1_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add introduction
            elements.append(Paragraph("Introduction", heading2_style))
            elements.append(Paragraph(doc_structure.get("introduction", self.description), normal_style))
            elements.append(Spacer(1, 0.3*inch))
            
            # Use our helper method to prepare deduplicated screenshots
            screenshot_data = self._prepare_screenshot_tracking()
            screenshot_map = screenshot_data["screenshot_map"]
            find_screenshot = screenshot_data["find_screenshot"]
            
            # Add each section with its content and screenshots
            for section_idx, section in enumerate(doc_structure.get("sections", [])):
                # Add section heading
                section_title = section.get("title", f"Section {section_idx+1}")
                elements.append(Paragraph(section_title, heading2_style))
                elements.append(Spacer(1, 0.1*inch))
                
                # Add section content
                section_content = section.get("content", "No content provided for this section.")
                if '**' in section_content:
                    formatted_content = self._format_content_for_pdf(section_content, normal_style, bold_style)
                    elements.extend(formatted_content)
                else:
                    elements.append(Paragraph(section_content, normal_style))
                elements.append(Spacer(1, 0.2*inch))
                
                # Add screenshots for this section
                if self.include_screenshots:
                    for ts in section.get("screenshot_timestamps", []):
                        if isinstance(ts, str):
                            try:
                                ts = float(ts)
                            except ValueError:
                                continue
                                
                        screenshot_result = find_screenshot(ts)
                        
                        if screenshot_result and screenshot_result[0] is not None:
                            img, reason = screenshot_result
                            
                            from ..utils.media_utils import format_timestamp
                            timestamp_text = f"Screenshot at {format_timestamp(ts)}"
                            elements.append(Paragraph(timestamp_text, 
                                                    ParagraphStyle('timestamp', 
                                                                parent=normal_style, 
                                                                fontName='Helvetica-Oblique')))
                            
                            img_buffer = BytesIO()
                            img.save(img_buffer, format="PNG")
                            img_buffer.seek(0)
                            
                            try:
                                img_for_pdf = Image(img_buffer, width=6*inch, height=3.5*inch)
                                elements.append(img_for_pdf)
                            except Exception as img_err:
                                logging.info(f"Error adding image to PDF: {img_err}")
                                elements.append(Paragraph(f"[Image placeholder - could not load image: {img_err}]", normal_style))
                            
                            elements.append(Spacer(1, 0.2*inch))
                    
                # Add process diagram if this is a process map section
                
                if section['title'] == 'Process Map' and section.get('diagram_image'):
                    try:
                        # Verify input image
                        if not isinstance(section['diagram_image'], PILImage.Image):
                            raise ValueError("diagram_image is not a valid PIL Image object")

                        # Calculate optimal dimensions
                        max_width_inches = 7.0
                        max_height_inches = 5.0
                        min_width_inches = 4.0  # Minimum width for readability
                        min_height_inches = 3.0  # Minimum height for readability
                        width_inches, height_inches = self._calculate_optimal_image_size_pdf(
                            section['diagram_image'], 
                            max_width_inches, 
                            max_height_inches,
                            min_width_inches,
                            min_height_inches
                        )

                        # Pre-resize the image using PIL
                        dpi = 150  # Higher DPI for better text clarity
                        width_pixels = int(width_inches * dpi)
                        height_pixels = int(height_inches * dpi)
                        resized_img = section['diagram_image'].resize(
                            (width_pixels, height_pixels),
                            PILImage.Resampling.LANCZOS  # High-quality resizing
                        )

                        # Convert resized image to bytes
                        diagram_buffer = BytesIO()
                        resized_img.convert('RGB').save(diagram_buffer, format='PNG', dpi=(dpi, dpi))
                        diagram_buffer.seek(0)

                        # Add diagram header
                        elements.append(Paragraph(
                            "Process Flow Diagram:",
                            ParagraphStyle(
                                name='diagram_header',
                                parent=normal_style,
                                fontSize=12,
                                fontName='Helvetica-Bold'
                            )
                        ))
                        elements.append(Spacer(1, 0.1*inch))

                        # Add the diagram image
                        diagram_img = Image(
                            diagram_buffer,
                            width=width_inches*inch,
                            height=height_inches*inch,
                            kind='proportional'
                        )
                        diagram_img.hAlign = 'CENTER'
                        elements.append(diagram_img)
                        elements.append(Spacer(1, 0.2*inch))

                        # Add mermaid code section if available
                        if section.get("mermaid_code"):
                            mermaid_code = section["mermaid_code"]
                            
                            # Add the mermaid code in a code block style
                            if mermaid_code:
                                elements.append(Paragraph(
                                    "Mermaid Diagram Code:",
                                    ParagraphStyle(
                                        name='code_header',
                                        parent=normal_style,
                                        fontSize=10,
                                        fontName='Helvetica-Bold'
                                    )
                                ))
                                elements.append(Spacer(1, 0.1 * inch))
                                elements.append(Preformatted(
                                    mermaid_code,
                                    ParagraphStyle(
                                        name='code_block',
                                        parent=normal_style,
                                        fontName='Courier',
                                        fontSize=8,
                                        leading=10,
                                        leftIndent=10,
                                        borderPadding=5,
                                    )
                                ))
                                elements.append(Spacer(1, 0.2 * inch))

                            elements.append(Spacer(1, 0.2*inch))
                            
                            # Generate editor URL with the mermaid code
                            mermaid_editor_url = self.generate_mermaid_editor_url(mermaid_code)
                            
                            # Add editor link
                            elements.append(Paragraph(
                                f'<b>Edit this diagram online:</b><br/>'
                                f'<a href="{mermaid_editor_url}" color="blue">Open in Mermaid Live Editor</a>',
                                ParagraphStyle(
                                    name='link_style',
                                    parent=normal_style,
                                    fontSize=10,
                                    leftIndent=10
                                )
                            ))
                            elements.append(Spacer(1, 0.1*inch))
                            
                            # Add instructions
                            elements.append(Paragraph(
                                '<i>Copy the mermaid code above and paste it into the editor to modify the diagram.</i>',
                                ParagraphStyle(
                                    name='instruction_style',
                                    parent=normal_style,
                                    fontSize=9,
                                    fontName='Helvetica-Oblique',
                                    leftIndent=10,
                                    textColor='#666666'
                                )
                            ))
                        else:
                            # Fallback if no mermaid code is available
                            elements.append(Paragraph(
                                f'<b>Edit diagrams online:</b><br/>'
                                f'<a href="https://mermaid.live/edit" color="blue">Mermaid Live Editor</a>',
                                ParagraphStyle(
                                    name='link_style',
                                    parent=normal_style,
                                    fontSize=10,
                                    leftIndent=10
                                )
                            ))
                        
                        elements.append(Spacer(1, 0.3*inch))
                        
                    except Exception as e:
                        logging.info(f"Error adding process diagram to PDF: {e}")
                        elements.append(Paragraph(
                            f"[Process diagram could not be loaded: {e}]",
                            normal_style
                        ))
                # Add subsections
                for subsection_idx, subsection in enumerate(section.get("subsections", [])):
                    subsection_title = subsection.get("title", f"Subsection {subsection_idx+1}")
                    elements.append(Paragraph(subsection_title, 
                                            ParagraphStyle('subsection', 
                                                        parent=heading2_style,
                                                        fontSize=14)))
                    elements.append(Spacer(1, 0.1*inch))
                    
                    subsection_content = subsection.get("content", "No content provided for this subsection.")
                    if '**' in subsection_content:
                        formatted_content = self._format_content_for_pdf(subsection_content, normal_style, bold_style)
                        elements.extend(formatted_content)
                    else:
                        elements.append(Paragraph(subsection_content, normal_style))
                    elements.append(Spacer(1, 0.15*inch))
                    
                    if self.include_screenshots:
                        for ts in subsection.get("screenshot_timestamps", []):
                            if isinstance(ts, str):
                                try:
                                    ts = float(ts)
                                except ValueError:
                                    continue
                                
                            screenshot_result = find_screenshot(ts)
                            
                            if screenshot_result and screenshot_result[0] is not None:
                                img, reason = screenshot_result
                                
                                from ..utils.media_utils import format_timestamp
                                timestamp_text = f"Screenshot at {format_timestamp(ts)}"
                                elements.append(Paragraph(timestamp_text, 
                                                        ParagraphStyle('timestamp', 
                                                                    parent=normal_style, 
                                                                    fontName='Helvetica-Oblique')))
                                
                                img_buffer = BytesIO()
                                img.save(img_buffer, format="PNG")
                                img_buffer.seek(0)
                                
                                try:
                                    img_for_pdf = Image(img_buffer, width=6*inch, height=3.5*inch)
                                    elements.append(img_for_pdf)
                                except Exception as img_err:
                                    logging.info(f"Error adding image to PDF: {img_err}")
                                    elements.append(Paragraph(f"[Image placeholder - could not load image: {img_err}]", normal_style))
                                
                                elements.append(Spacer(1, 0.2*inch))
                    
                if section_idx < len(doc_structure.get("sections", [])) - 1:
                    elements.append(Spacer(1, 0.3*inch))
            
            # Build the PDF in memory
            doc.build(elements)
            buffer.seek(0)
            return buffer.getvalue()
        
        else:
            from docx import Document
            doc_buffer = BytesIO()
            doc = Document()
            
            # Set document title and add heading
            doc.add_heading(self.title, 0)
            
            # Add document description
            doc.add_paragraph(self.description)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add a table of contents header
            doc.add_heading('Contents', level=1)
            if self.include_screenshots:
                doc.add_paragraph('Screenshots and Descriptions')
            
            # Generate descriptions for each screenshot
            descriptions = self._generate_section_descriptions()
            
            # Add document title
            doc.add_heading(doc_structure.get("title", self.title), level=1)
            
            # Add introduction
            doc.add_heading("Introduction", level=2)
            doc.add_paragraph(doc_structure.get("introduction", self.description))
            
            # Use our helper method to prepare deduplicated screenshots
            screenshot_data = self._prepare_screenshot_tracking()
            screenshot_map = screenshot_data["screenshot_map"]
            find_screenshot = screenshot_data["find_screenshot"]
            
            # Add each section with its content and screenshots
            for section_idx, section in enumerate(doc_structure.get("sections", [])):
                # Add section heading
                section_title = section.get("title", f"Section {section_idx+1}")
                doc.add_heading(section_title, level=2)
                
                # Add section content
                section_content = section.get("content", "No content provided for this section.")
                if '**' in section_content:
                    lines = section_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        if line.startswith('**') and line.endswith('**'):
                            heading_text = line.strip('*')
                            para = doc.add_paragraph()
                            para.add_run(heading_text).bold = True
                        elif line.startswith('• ') or line.startswith('- '):
                            bullet_text = line[2:]
                            doc.add_paragraph(bullet_text, style='List Bullet')
                        elif line and line[0].isdigit() and '. ' in line:
                            doc.add_paragraph(line, style='List Number')
                        else:
                            doc.add_paragraph(line)
                else:
                    paragraphs = section_content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                
                # Add screenshots for this section
                if self.include_screenshots:
                    for ts in section.get("screenshot_timestamps", []):
                        if isinstance(ts, str):
                            try:
                                ts = float(ts)
                            except ValueError:
                                continue
                                
                        screenshot_result = find_screenshot(ts)
                        
                        if screenshot_result and screenshot_result[0] is not None:
                            img, reason = screenshot_result
                            
                            from ..utils.media_utils import format_timestamp
                            timestamp_text = f"Screenshot at {format_timestamp(ts)}"
                            timestamp_para = doc.add_paragraph()
                            timestamp_para.add_run(timestamp_text).italic = True
                            
                            img_buffer = BytesIO()
                            img.save(img_buffer, format="PNG")
                            img_buffer.seek(0)
                            
                            try:
                                doc.add_picture(img_buffer, width=Inches(6))
                            except Exception as img_err:
                                logging.info(f"Error adding image to DOCX: {img_err}")
                                doc.add_paragraph(f"[Image placeholder - could not load image: {img_err}]")
                    
                # Add process diagram if this is a process map section
                if section['title'] == 'Process Map' and section.get('diagram_image'):
                    try:
                        temp_path = "temp_process_diagram.png"
                        section['diagram_image'].save(temp_path, format='PNG')
                        optimal_width, optimal_height = self._calculate_optimal_image_size(section['diagram_image'])

                        # Add diagram header
                        doc.add_paragraph("Process Flow Diagram:", style='Heading 3')
                        
                        # Add the diagram image
                        doc.add_picture(temp_path, width=Inches(optimal_width))
                        logging.info(f"DEBUG: Added process diagram to DOCX: {temp_path}")
                        
                        # Add some spacing
                        doc.add_paragraph()
                        
                        # Add mermaid code section if available
                        if section.get("mermaid_code"):
                            mermaid_code = section["mermaid_code"]
                            
                            # Add mermaid code header
                            mermaid_header = doc.add_paragraph("Mermaid Diagram Code:", style='Heading 4')
                            
                            # Add the mermaid code in a code block style
                            code_paragraph = doc.add_paragraph()
                            code_run = code_paragraph.add_run(mermaid_code)
                            
                            # Style the code block
                            code_run.font.name = 'Courier New'
                            code_run.font.size = Inches(0.08)  # Approximately 9pt
                            
                            # Add background color to the paragraph (light gray)
                            from docx.oxml.ns import qn
                            pPr = code_paragraph._element.get_or_add_pPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), 'F5F5F5')  # Light gray background
                            pPr.append(shd)
                            
                            # Add some spacing
                            doc.add_paragraph()
                            
                            # Generate editor URL with the mermaid code
                            mermaid_editor_url = self.generate_mermaid_editor_url_docx(mermaid_code)
                            
                            # Add editor link section
                            link_paragraph = doc.add_paragraph()
                            link_paragraph.add_run("Edit this diagram online: ")
                            
                            # Try to add hyperlink, fallback to plain text if it fails
                            if not self.add_hyperlink(link_paragraph, "Open in Mermaid Live Editor", mermaid_editor_url):
                                link_paragraph.add_run(f"Open in Mermaid Live Editor ({mermaid_editor_url})")
                            
                            # Add instructions
                            instruction_paragraph = doc.add_paragraph()
                            instruction_run = instruction_paragraph.add_run(
                                "Copy the mermaid code above and paste it into the editor to modify the diagram."
                            )
                            instruction_run.italic = True
                            instruction_run.font.size = Inches(0.09)  # Approximately 8pt
                            
                        else:
                            # Fallback if no mermaid code is available
                            link_paragraph = doc.add_paragraph()
                            link_paragraph.add_run("Edit diagrams online: ")
                            
                            if not self.add_hyperlink(link_paragraph, "Mermaid Live Editor", "https://mermaid.live/edit"):
                                link_paragraph.add_run("Mermaid Live Editor (https://mermaid.live/edit)")
                        
                        # Add final spacing
                        doc.add_paragraph()
                        
                        # Clean up temp file
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                            
                    except Exception as e:
                        logging.info(f"Error adding process diagram to DOCX: {e}")
                        doc.add_paragraph(f"[Process diagram could not be loaded: {e}]")
                # Add subsections
                for subsection_idx, subsection in enumerate(section.get("subsections", [])):
                    subsection_title = subsection.get("title", f"Subsection {subsection_idx+1}")
                    doc.add_heading(subsection_title, level=3)
                    
                    subsection_content = subsection.get("content", "No content provided for this subsection.")
                    if '**' in subsection_content:
                        lines = subsection_content.split('\n')
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue
                            if line.startswith('**') and line.endswith('**'):
                                heading_text = line.strip('*')
                                para = doc.add_paragraph()
                                para.add_run(heading_text).bold = True
                            elif line.startswith('• ') or line.startswith('- '):
                                bullet_text = line[2:]
                                doc.add_paragraph(bullet_text, style='List Bullet')
                            elif line and line[0].isdigit() and '. ' in line:
                                doc.add_paragraph(line, style='List Number')
                            else:
                                doc.add_paragraph(line)
                    else:
                        paragraphs = subsection_content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para.strip())
                    
                    if self.include_screenshots:
                        for ts in subsection.get("screenshot_timestamps", []):
                            if isinstance(ts, str):
                                try:
                                    ts = float(ts)
                                except ValueError:
                                    continue
                                
                            screenshot_result = find_screenshot(ts)
                            
                            if screenshot_result and screenshot_result[0] is not None:
                                img, reason = screenshot_result
                                
                                from ..utils.media_utils import format_timestamp
                                timestamp_text = f"Screenshot at {format_timestamp(ts)}"
                                timestamp_para = doc.add_paragraph()
                                timestamp_para.add_run(timestamp_text).italic = True
                                
                                img_buffer = BytesIO()
                                img.save(img_buffer, format="PNG")
                                img_buffer.seek(0)
                                
                                try:
                                    doc.add_picture(img_buffer, width=Inches(6))
                                except Exception as img_err:
                                    logging.info(f"Error adding image to DOCX: {img_err}")
                                    doc.add_paragraph(f"[Image placeholder - could not load image: {img_err}]")
            
            # Save the document to the buffer
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            return doc_buffer.getvalue()